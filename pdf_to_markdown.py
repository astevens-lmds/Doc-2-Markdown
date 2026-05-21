"""
PDF to Markdown Converter
A comprehensive desktop application to convert PDF files to Markdown format.
Features: AI enhancement, OCR, table detection, multiple export formats, drag & drop.
"""

__version__ = "2.3.5"
__version_date__ = "2026-05-12"

import threading
import queue
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path
import tkinter as tk
from tkinter import filedialog, messagebox, ttk, scrolledtext, simpledialog
from datetime import datetime
import os
import sys
import json
import subprocess
import platform

# Monkey patch subprocess.Popen to suppress console window on Windows
if platform.system() == "Windows":
    _original_Popen = subprocess.Popen

    class NoConsolePopen(_original_Popen):
        def __init__(self, args, **kwargs):
            if 'startupinfo' not in kwargs:
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW
                startupinfo.wShowWindow = subprocess.SW_HIDE
                kwargs['startupinfo'] = startupinfo
            super().__init__(args, **kwargs)
    
    subprocess.Popen = NoConsolePopen
import urllib.request
import urllib.error
import base64
import re
import tempfile
import subprocess
from concurrent.futures import ThreadPoolExecutor, as_completed
import queue

# Configuration and cost tracking files.
# When running from a read-only location (e.g. a mounted .dmg app bundle),
# the launcher sets DOC2MD_DATA_DIR to a writable path; otherwise fall back
# to the source directory (dev mode).
_DATA_DIR = Path(os.environ.get("DOC2MD_DATA_DIR") or Path(__file__).parent)
_DATA_DIR.mkdir(parents=True, exist_ok=True)
CONFIG_FILE = _DATA_DIR / "config.json"
USAGE_FILE = _DATA_DIR / "usage.json"
PROMPTS_FILE = _DATA_DIR / "custom_prompts.json"

# ============================================================================
# LAZY IMPORT SYSTEM - Performance Optimization
# Heavy dependencies are imported only when needed to speed up app startup
# ============================================================================

# Availability flags (checked at startup, imports deferred)
HAS_PDFPLUMBER = False
HAS_TESSERACT = False
HAS_DOCX = False
HAS_MSG = False
HAS_DND = False
HAS_TIKA = False
HAS_TABULA = False
HAS_PYMUPDF4LLM = False
HAS_TIKTOKEN = False
HAS_EPUB = False
HAS_MOBI = False
HAS_RAG = False
HAS_CHROMADB = False
HAS_NUMPY = False

# Lazy-loaded module cache
_lazy_modules = {}

def _lazy_import(module_name, package=None):
    """Lazy import a module only when first accessed."""
    cache_key = f"{package}.{module_name}" if package else module_name
    if cache_key not in _lazy_modules:
        import importlib
        try:
            if package:
                _lazy_modules[cache_key] = importlib.import_module(module_name, package)
            else:
                _lazy_modules[cache_key] = importlib.import_module(module_name)
        except ImportError:
            _lazy_modules[cache_key] = None
    return _lazy_modules[cache_key]

# Quick availability checks (fast - no heavy imports)
def _check_availability():
    """Quick check which optional packages are available."""
    global HAS_PDFPLUMBER, HAS_TESSERACT, HAS_DOCX, HAS_MSG, HAS_DND
    global HAS_TIKA, HAS_TABULA, HAS_PYMUPDF4LLM, HAS_TIKTOKEN
    global HAS_EPUB, HAS_MOBI, HAS_RAG, HAS_CHROMADB, HAS_NUMPY

    import importlib.util

    HAS_PDFPLUMBER = importlib.util.find_spec("pdfplumber") is not None
    HAS_TESSERACT = importlib.util.find_spec("pytesseract") is not None
    HAS_DOCX = importlib.util.find_spec("docx") is not None
    HAS_MSG = importlib.util.find_spec("extract_msg") is not None
    try:
        HAS_DND = importlib.util.find_spec("tkinterdnd2") is not None
    except (ValueError, ModuleNotFoundError):
        HAS_DND = False
    HAS_TIKA = importlib.util.find_spec("tika") is not None
    HAS_TABULA = importlib.util.find_spec("tabula") is not None
    HAS_PYMUPDF4LLM = importlib.util.find_spec("pymupdf4llm") is not None
    HAS_TIKTOKEN = importlib.util.find_spec("tiktoken") is not None
    HAS_EPUB = importlib.util.find_spec("ebooklib") is not None
    HAS_MOBI = importlib.util.find_spec("mobi") is not None
    HAS_NUMPY = importlib.util.find_spec("numpy") is not None

    # Check RAG module availability
    try:
        from rag_module import HAS_CHROMADB as RAG_HAS_CHROMADB, HAS_NUMPY as RAG_HAS_NUMPY
        HAS_RAG = True
        HAS_CHROMADB = RAG_HAS_CHROMADB
        HAS_NUMPY = RAG_HAS_NUMPY
    except ImportError:
        HAS_RAG = False

# Run availability check at startup (fast)
_check_availability()

# Lazy import functions for heavy modules (called only when needed)
def get_pdfplumber():
    """Get pdfplumber module (lazy loaded)."""
    return _lazy_import("pdfplumber") if HAS_PDFPLUMBER else None

def get_pytesseract():
    """Get pytesseract module (lazy loaded)."""
    return _lazy_import("pytesseract") if HAS_TESSERACT else None

def get_pil_image():
    """Get PIL.Image module (lazy loaded)."""
    if HAS_TESSERACT:
        pil = _lazy_import("PIL")
        return pil.Image if pil else None
    return None

def get_docx():
    """Get python-docx Document class (lazy loaded)."""
    if HAS_DOCX:
        docx = _lazy_import("docx")
        return docx.Document if docx else None
    return None

def get_extract_msg():
    """Get extract_msg module (lazy loaded)."""
    return _lazy_import("extract_msg") if HAS_MSG else None

def get_tika_parser():
    """Get tika parser (lazy loaded)."""
    if HAS_TIKA:
        tika = _lazy_import("tika")
        if tika:
            from tika import parser
            return parser
    return None

def get_tabula():
    """Get tabula module (lazy loaded)."""
    return _lazy_import("tabula") if HAS_TABULA else None

def get_pymupdf4llm():
    """Get pymupdf4llm module (lazy loaded)."""
    return _lazy_import("pymupdf4llm") if HAS_PYMUPDF4LLM else None

def get_tiktoken():
    """Get tiktoken module (lazy loaded)."""
    return _lazy_import("tiktoken") if HAS_TIKTOKEN else None

def get_ebooklib():
    """Get ebooklib module (lazy loaded)."""
    return _lazy_import("ebooklib") if HAS_EPUB else None

def get_mobi():
    """Get mobi module (lazy loaded)."""
    return _lazy_import("mobi") if HAS_MOBI else None

def get_rag_module():
    """Get RAG module components (lazy loaded)."""
    if HAS_RAG:
        try:
            from rag_module import (
                RAGChunker, EmbeddingClient, VectorStore, SimpleVectorStore,
                HybridRetriever, get_vector_store, vectorize_markdown_file
            )
            return {
                'RAGChunker': RAGChunker,
                'EmbeddingClient': EmbeddingClient,
                'VectorStore': VectorStore,
                'SimpleVectorStore': SimpleVectorStore,
                'HybridRetriever': HybridRetriever,
                'get_vector_store': get_vector_store,
                'vectorize_markdown_file': vectorize_markdown_file
            }
        except ImportError:
            pass
    return None

# For backward compatibility - these will be set when modules are first used
pdfplumber = None
pytesseract = None
Image = None
Document = None
extract_msg = None
tiktoken = None
ebooklib = None
epub = None
mobi = None

# TkinterDnD must be imported at module level for GUI startup
TkinterDnD = None
DND_FILES = None
if HAS_DND:
    try:
        from tkinterdnd2 import DND_FILES, TkinterDnD
    except ImportError:
        HAS_DND = False

import email
from email import policy
from email.parser import BytesParser


def check_dependencies():
    """Check required dependencies."""
    required = {'pymupdf': 'PyMuPDF', 'PIL': 'Pillow'}
    if not HAS_DOCX:
        required['python-docx'] = 'python-docx'
    if not HAS_MSG:
        required['extract-msg'] = 'extract-msg'
    if not HAS_TIKTOKEN:
        required['tiktoken'] = 'tiktoken'
    if not HAS_EPUB:
        required['ebooklib'] = 'ebooklib'
    if not HAS_MOBI:
        required['mobi'] = 'mobi'
        
    missing = []
    for module, package in required.items():
        try:
            if module == 'pymupdf':
                import fitz
            elif module == 'python-docx':
                import docx
            elif module == 'extract-msg':
                import extract_msg
            elif module == 'PIL':
                from PIL import Image
        except ImportError:
            missing.append(package)
    return missing


def setup_java_environment(config):
    """Setup Java environment for Tika and Tabula."""
    java_path = config.get("java_path", "")
    if java_path and os.path.exists(java_path):
        os.environ["JAVA_HOME"] = str(Path(java_path).parent.parent)
        os.environ["PATH"] = str(Path(java_path).parent) + os.pathsep + os.environ["PATH"]


def load_config():
    """Load configuration from file."""
    default_config = {
        "monthly_budget": 30.00,
        "use_ai_enhancement": True,
        "use_ocr": False,
        "use_table_detection": True,
        "parallel_workers": 2,
        "skip_existing": False,
        "output_format": "markdown",
        "output_template": "{name}",
        "include_metadata": True,
        "custom_prompt": "",
        "vision_model": "openai/gpt-4o-mini",  # Default vision model
        
        # New provider structure
        "active_provider": "datalab",
        "default_model": "datalab/marker-ocr",
        "api_keys": {
            "openrouter": "",
            "openai": "",
            "anthropic": "",
            "google": "",
            "datalab": ""
        },
        "use_tika": False,
        "use_tabula": True,
        "java_path": "",  # Set to your java.exe path if using Tika/Tabula (e.g., "C:\\Program Files\\Java\\jdk-17\\bin\\java.exe")

        # RAG Settings
        "rag_settings": {
            "enabled": True,
            "auto_vectorize": True,
            "embedding_provider": "openai",
            "embedding_model": "text-embedding-3-large",
            "vector_db_path": "./vector_db",
            "default_doc_type": "default",
            "chunk_overlap_pct": 0.12,
            "retrieval_k": 10,
            "hybrid_search": True
        }
    }
    
    if CONFIG_FILE.exists():
        try:
            with open(CONFIG_FILE, 'r') as f:
                config = json.load(f)
                
                # Migration: Move old 'openrouter_api_key' to new structure
                if "openrouter_api_key" in config and not config.get("api_keys"):
                    config["api_keys"] = default_config["api_keys"].copy()
                    config["api_keys"]["openrouter"] = config.pop("openrouter_api_key")
                
                # Merge defaults
                for key, value in default_config.items():
                    if key not in config:
                        config[key] = value
                    elif isinstance(value, dict) and isinstance(config[key], dict):
                        for k, v in value.items():
                            if k not in config[key]:
                                config[key][k] = v
                                
                return config
        except Exception:
            pass
    return default_config


def save_config(config):
    """Save configuration to file."""
    with open(CONFIG_FILE, 'w') as f:
        json.dump(config, f, indent=4)


def load_usage():
    """Load usage tracking data."""
    default_usage = {
        "month": datetime.now().strftime("%Y-%m"),
        "total_cost": 0.0,
        "total_tokens": 0,
        "conversions": 0
    }
    if USAGE_FILE.exists():
        try:
            with open(USAGE_FILE, 'r') as f:
                usage = json.load(f)
                current_month = datetime.now().strftime("%Y-%m")
                if usage.get("month") != current_month:
                    usage = default_usage
                return usage
        except Exception:
            pass
    return default_usage


def save_usage(usage):
    """Save usage tracking data."""
    with open(USAGE_FILE, 'w') as f:
        json.dump(usage, f, indent=4)


# Available models with pricing (per 1M tokens)
# Provider Settings
PROVIDERS = {
    "openrouter": {
        "name": "OpenRouter (Aggregator)",
        "url": "https://openrouter.ai/keys",
        "base_url": "https://openrouter.ai/api/v1/chat/completions"
    },
    "openai": {
        "name": "OpenAI",
        "url": "https://platform.openai.com/api-keys",
        "base_url": "https://api.openai.com/v1/chat/completions"
    },
    "anthropic": {
        "name": "Anthropic",
        "url": "https://console.anthropic.com/settings/keys",
        "base_url": "https://api.anthropic.com/v1/messages"
    },
    "google": {
        "name": "Google Gemini",
        "url": "https://aistudio.google.com/app/apikey",
        "base_url": "https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent"
    },
    "datalab": {
        "name": "DataLabs OCR",
        "url": "https://datalab.to",
        "base_url": "https://www.datalab.to/api/v1/marker"
    }
}

# Models by Provider (Updated April 2026 - Verified Available)
PROVIDER_MODELS = {
    "openrouter": {
        # === GOOGLE GEMINI (Most Reliable) ===
        "google/gemini-3.1-pro": {"name": "Gemini 3.1 Pro (Recommended)", "input": 0.50, "output": 1.50, "vision": True},
        "google/gemini-3.0-flash": {"name": "Gemini 3.0 Flash (Fast)", "input": 0.05, "output": 0.20, "vision": True},
        "google/gemini-3.0-flash-lite": {"name": "Gemini 3.0 Flash Lite", "input": 0.02, "output": 0.08, "vision": True},
        # === ANTHROPIC CLAUDE ===
        "anthropic/claude-opus-4-7": {"name": "Claude Opus 4.7 (Most Capable)", "input": 15.00, "output": 45.00, "vision": True},
        "anthropic/claude-sonnet-4-6": {"name": "Claude Sonnet 4.6", "input": 3.00, "output": 15.00, "vision": True},
        "anthropic/claude-haiku-4-5": {"name": "Claude Haiku 4.5", "input": 0.25, "output": 1.25, "vision": True},
        # === OPENAI GPT ===
        "openai/gpt-6-omni": {"name": "GPT-6 Omni", "input": 1.50, "output": 5.00, "vision": True},
        "openai/gpt-5.5-mini": {"name": "GPT-5.5 Mini", "input": 0.10, "output": 0.30, "vision": True},
        # === META LLAMA ===
        "meta-llama/llama-4-100b-instruct": {"name": "Llama 4 100B", "input": 0.10, "output": 0.25, "vision": True},
        # === DEEPSEEK ===
        "deepseek/deepseek-v4-chat": {"name": "DeepSeek V4 Chat", "input": 0.10, "output": 0.20, "vision": True},
        "deepseek/deepseek-r2:free": {"name": "DeepSeek R2 (FREE)", "input": 0.00, "output": 0.00, "vision": False},
    },
    "openai": {
        "gpt-6-omni": {"name": "GPT-6 Omni", "input": 1.50, "output": 5.00, "vision": True},
        "gpt-6-omni-mini": {"name": "GPT-6 Omni Mini", "input": 0.10, "output": 0.30, "vision": True},
        "gpt-5.5-mini": {"name": "GPT-5.5 Mini", "input": 0.10, "output": 0.30, "vision": True},
    },
    "anthropic": {
        "claude-opus-4-7": {"name": "Claude Opus 4.7", "input": 15.00, "output": 45.00, "vision": True},
        "claude-sonnet-4-6": {"name": "Claude Sonnet 4.6 (Recommended)", "input": 3.00, "output": 15.00, "vision": True},
        "claude-haiku-4-5": {"name": "Claude Haiku 4.5", "input": 0.25, "output": 1.25, "vision": True},
    },
    "google": {
        "gemini-3.1-pro": {"name": "Gemini 3.1 Pro", "input": 0.50, "output": 1.50, "vision": True},
        "gemini-3.0-flash": {"name": "Gemini 3.0 Flash", "input": 0.05, "output": 0.20, "vision": True},
        "gemini-3.0-flash-lite": {"name": "Gemini 3.0 Flash Lite", "input": 0.02, "output": 0.08, "vision": True},
    },
    "datalab": {
        "datalab/marker-ocr": {"name": "DataLabs Marker OCR (High-End)", "input": 0.0, "output": 0.0, "vision": True}
    }
}

# Flatten for easy lookup
MODELS = {k: v for p in PROVIDER_MODELS.values() for k, v in p.items()}

# Flatten for easy lookup
MODELS = {k: v for p in PROVIDER_MODELS.values() for k, v in p.items()}

DEFAULT_PROMPT = """You are a PDF text formatter. Your ONLY job is to clean up and format the extracted text below into Markdown.

CRITICAL RULES:
1. ONLY output the actual content from the text provided below
2. If the text is empty, garbled, or unreadable, output exactly: [NO READABLE TEXT EXTRACTED]
3. NEVER create example content, tutorials, or placeholder text
4. NEVER invent or hallucinate content that is not in the source text
5. Do NOT add introductions, explanations, or meta-commentary

Formatting Instructions (LEGAL-RAG SPECIFIC):
- Fix formatting issues or artifacts from PDF extraction
- **Hierarchy:** Use # for Document Title, ## for Main Headings (e.g., Arguments, Sections), ### for Sub-headings.
- **Tables:** Convert tables into proper Markdown table format.
- **Lists:** Format lists (bulleted and numbered) correctly.
- **Citations:** Preserve all citations exactly as they appear.
- **Line Breaks:** Clean up hyphenation from line breaks but PRESERVE paragraph structure carefully.

If you cannot find real document content in the text below, respond ONLY with: [NO READABLE TEXT EXTRACTED]

Text to convert:
---
{text}
---

Output the formatted document content (or [NO READABLE TEXT EXTRACTED] if empty/unreadable):"""


class OpenRouterClient:
    """Client for OpenRouter API with vision support."""

    def __init__(self, api_key):
        self.api_key = api_key
        self.base_url = "https://openrouter.ai/api/v1/chat/completions"

    def enhance_markdown(self, raw_text, model="anthropic/claude-sonnet-4", custom_prompt=None):
        """Send text to AI for markdown enhancement."""
        prompt_template = custom_prompt if custom_prompt else DEFAULT_PROMPT

        data = {
            "model": model,
            "messages": [{"role": "user", "content": prompt_template.format(text=raw_text[:15000])}],
            "max_tokens": 8000,
            "temperature": 0.1
        }

        return self._make_request(data)

    def process_with_vision(self, image_base64, model="openai/gpt-4o-mini", prompt=None):
        """Process image with vision model for better extraction."""
        if prompt is None:
            prompt = "Extract all text from this image and format it as clean Markdown. Preserve tables, lists, and formatting."

        data = {
            "model": model,
            "messages": [{
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{image_base64}"}}
                ]
            }],
            "max_tokens": 4000,
            "temperature": 0.1
        }

        return self._make_request(data)

    # Max time we're willing to sit and wait on a single 429 backoff.
    # Free-tier OpenRouter resets can be hours away; past this we bail with
    # a clear error so the user can switch models instead of hanging.
    _RATE_LIMIT_MAX_WAIT = 90

    @staticmethod
    def _parse_retry_after(headers, body):
        """Return seconds to wait for a 429, or None if no signal."""
        import time
        # Standard HTTP header
        retry_after = headers.get('Retry-After')
        if retry_after:
            try:
                return max(0, int(float(retry_after)))
            except (TypeError, ValueError):
                pass
        # OpenRouter: X-RateLimit-Reset is ms-epoch
        reset = headers.get('X-RateLimit-Reset')
        if reset:
            try:
                reset_ms = int(reset)
                wait = (reset_ms / 1000.0) - time.time()
                return max(0, int(wait))
            except (TypeError, ValueError):
                pass
        # Sometimes the body carries the reset too
        try:
            parsed = json.loads(body) if body else {}
            meta = parsed.get('error', {}).get('metadata', {})
            inner = meta.get('headers', {})
            reset_ms = inner.get('X-RateLimit-Reset')
            if reset_ms:
                wait = (int(reset_ms) / 1000.0) - time.time()
                return max(0, int(wait))
        except (ValueError, TypeError, AttributeError):
            pass
        return None

    def _make_request(self, data, max_retries=3):
        """Make API request to OpenRouter with retry logic."""
        import time
        from http.client import IncompleteRead

        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json",
            "HTTP-Referer": "https://github.com/pdf-to-markdown",
            "X-Title": "PDF to Markdown Converter"
        }

        last_error = None
        for attempt in range(max_retries):
            req = urllib.request.Request(
                self.base_url,
                data=json.dumps(data).encode('utf-8'),
                headers=headers,
                method='POST'
            )

            try:
                with urllib.request.urlopen(req, timeout=300) as response:
                    result = json.loads(response.read().decode('utf-8'))
                    content = result['choices'][0]['message']['content']
                    usage = result.get('usage', {})
                    return {
                        'content': content,
                        'input_tokens': usage.get('prompt_tokens', 0),
                        'output_tokens': usage.get('completion_tokens', 0),
                        'model': data['model']
                    }
            except IncompleteRead as e:
                last_error = f"Incomplete response (attempt {attempt+1}/{max_retries}): {e}"
                if attempt < max_retries - 1:
                    time.sleep(2 ** attempt)  # Exponential backoff
                    continue
            except urllib.error.HTTPError as e:
                error_body = e.read().decode('utf-8') if e.fp else str(e)
                # Honor 429 with header-aware backoff. If the reset window is
                # further out than _RATE_LIMIT_MAX_WAIT, bail immediately with
                # a clear message — the caller should pick a different model.
                if e.code == 429:
                    wait = self._parse_retry_after(e.headers, error_body)
                    if wait is not None and wait <= self._RATE_LIMIT_MAX_WAIT and attempt < max_retries - 1:
                        # Add a small jitter floor so we don't hammer the moment it opens
                        time.sleep(max(2, wait) + 1)
                        last_error = f"Rate limited (429), waited {wait}s and retrying"
                        continue
                    model = data.get('model', '<unknown>')
                    hint = (
                        f"Daily/minute quota exhausted for model '{model}' on OpenRouter."
                        " Switch to a different model in Settings (e.g. anthropic/claude-haiku-4-5"
                        " or google/gemini-2.5-flash) or wait for the reset window."
                    )
                    if wait is not None:
                        hint += f" Reset in ~{wait}s."
                    raise Exception(f"API Error (429): {hint} Raw: {error_body}")
                # Retry on 5xx errors
                if e.code >= 500 and attempt < max_retries - 1:
                    last_error = f"Server error ({e.code}), retrying..."
                    time.sleep(2 ** attempt)
                    continue
                raise Exception(f"API Error ({e.code}): {error_body}")
            except urllib.error.URLError as e:
                last_error = f"Connection error (attempt {attempt+1}/{max_retries}): {e.reason}"
                if attempt < max_retries - 1:
                    time.sleep(2 ** attempt)
                    continue
            except TimeoutError:
                last_error = f"Request timeout (attempt {attempt+1}/{max_retries})"
                if attempt < max_retries - 1:
                    time.sleep(2 ** attempt)
                    continue

        raise Exception(f"Request failed after {max_retries} attempts: {last_error}")

    def estimate_cost(self, text_length, model):
        """Estimate cost for processing text."""
        # Rough estimate: 4 chars per token
        est_input_tokens = text_length / 4
        est_output_tokens = est_input_tokens * 0.8
        return self.calculate_cost(int(est_input_tokens), int(est_output_tokens), model)

    def calculate_cost(self, input_tokens, output_tokens, model):
        """Calculate cost based on token usage."""
        if model not in MODELS:
            return 0.0
        pricing = MODELS[model]
        # Support both 'input'/'output' and 'input_cost'/'output_cost' for legacy
        i_cost = pricing.get('input', pricing.get('input_cost', 0))
        o_cost = pricing.get('output', pricing.get('output_cost', 0))
        return (input_tokens / 1_000_000) * i_cost + (output_tokens / 1_000_000) * o_cost


class OpenAIClient(OpenRouterClient):
    """Client for OpenAI API."""
    def __init__(self, api_key):
        super().__init__(api_key)
        self.base_url = "https://api.openai.com/v1/chat/completions"
        
    def _make_request(self, data):
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
        req = urllib.request.Request(
            self.base_url,
            data=json.dumps(data).encode('utf-8'),
            headers=headers,
            method='POST'
        )
        try:
            with urllib.request.urlopen(req) as response:
                result = json.loads(response.read().decode('utf-8'))
                return {
                    'content': result['choices'][0]['message']['content'],
                    'input_tokens': result.get('usage', {}).get('prompt_tokens', 0),
                    'output_tokens': result.get('usage', {}).get('completion_tokens', 0),
                    'model': data['model']
                }
        except Exception as e:
            raise Exception(f"OpenAI Error: {e}")


class AnthropicClient(OpenRouterClient):
    """Client for Anthropic API."""
    def __init__(self, api_key):
        super().__init__(api_key)
        self.base_url = "https://api.anthropic.com/v1/messages"
        
    def enhance_markdown(self, raw_text, model="claude-3-5-sonnet-20241022", custom_prompt=None):
        prompt_template = custom_prompt if custom_prompt else DEFAULT_PROMPT
        data = {
            "model": model,
            "messages": [{"role": "user", "content": prompt_template.format(text=raw_text[:15000])}],
            "max_tokens": 4096,
            "temperature": 0.1
        }
        return self._make_request(data)

    def process_with_vision(self, image_base64, model="claude-3-5-sonnet-20241022", prompt=None):
        if prompt is None:
            prompt = "Extract all text from this image and format it as clean Markdown."
            
        data = {
            "model": model,
            "messages": [{
                "role": "user",
                "content": [
                    {"type": "image", "source": {"type": "base64", "media_type": "image/jpeg", "data": image_base64}},
                    {"type": "text", "text": prompt}
                ]
            }],
            "max_tokens": 4096,
            "temperature": 0.1
        }
        return self._make_request(data)

    def _make_request(self, data):
        headers = {
            "x-api-key": self.api_key,
            "anthropic-version": "2023-06-01",
            "content-type": "application/json"
        }
        req = urllib.request.Request(
            self.base_url,
            data=json.dumps(data).encode('utf-8'),
            headers=headers,
            method='POST'
        )
        with urllib.request.urlopen(req) as response:
            resp_json = json.loads(response.read().decode('utf-8'))
            usage = resp_json.get('usage', {})
            return {
                'content': resp_json['content'][0]['text'],
                'input_tokens': usage.get('input_tokens', 0),
                'output_tokens': usage.get('output_tokens', 0),
                'model': data['model']
            }


class GoogleClient(OpenRouterClient):
    """Client for Google Gemini API."""
    def __init__(self, api_key):
        super().__init__(api_key)
        self.base_url_template = "https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={key}"

    def enhance_markdown(self, raw_text, model="gemini-1.5-flash", custom_prompt=None):
        prompt_template = custom_prompt if custom_prompt else DEFAULT_PROMPT
        url = self.base_url_template.format(model=model, key=self.api_key)
        
        data = {
            "contents": [{
                "parts": [{"text": prompt_template.format(text=raw_text[:15000])}]
            }],
            "generationConfig": {
                "temperature": 0.1,
                "maxOutputTokens": 8000
            }
        }
        return self._make_request(url, data, model)

    def process_with_vision(self, image_base64, model="gemini-1.5-flash", prompt=None):
        if prompt is None:
            prompt = "Extract text."
        url = self.base_url_template.format(model=model, key=self.api_key)
        
        data = {
            "contents": [{
                "parts": [
                    {"text": prompt},
                    {"inline_data": {"mime_type": "image/jpeg", "data": image_base64}}
                ]
            }]
        }
        return self._make_request(url, data, model)

    def _make_request(self, url, data, model):
        headers = {"Content-Type": "application/json"}
        req = urllib.request.Request(
            url,
            data=json.dumps(data).encode('utf-8'),
            headers=headers,
            method='POST'
        )
        with urllib.request.urlopen(req) as response:
            resp_json = json.loads(response.read().decode('utf-8'))
            try:
                content = resp_json['candidates'][0]['content']['parts'][0]['text']
                usage_meta = resp_json.get('usageMetadata', {})
                return {
                    'content': content,
                    'input_tokens': usage_meta.get('promptTokenCount', 0),
                    'output_tokens': usage_meta.get('candidatesTokenCount', 0),
                    'model': model
                }
            except:
                raise Exception(f"Google Error: {resp_json}")


class ClientFactory:
    @staticmethod
    def get_client(provider, api_key):
        if provider == "openai":
            return OpenAIClient(api_key)
        elif provider == "anthropic":
            return AnthropicClient(api_key)
        elif provider == "google":
            return GoogleClient(api_key)
        elif provider == "datalab":
            return DatalabClient(api_key)
        else:
            return OpenRouterClient(api_key)

import requests
import time

class DatalabClient(OpenRouterClient):
    """Client for DataLabs OCR API for High-End Conversion."""
    def __init__(self, api_key):
        super().__init__(api_key)
        self.base_url = "https://www.datalab.to/api/v1/marker"

    def process_document(self, file_path):
        headers = {"X-Api-Key": self.api_key}
        with open(file_path, 'rb') as f:
            files = {'file': (os.path.basename(file_path), f, 'application/pdf')}
            data = {'langs': 'English', 'force_ocr': 'false', 'paginate': 'true'}
            req = requests.post(self.base_url, headers=headers, files=files, data=data)

        if req.status_code != 200:
            raise Exception(f"Datalab API Error: {req.text}")
            
        result = req.json()
        request_check_url = result.get('request_check_url')
        
        if not request_check_url:
            # Maybe it returns markdown synchronously
            if 'markdown' in result:
                return {'content': result['markdown'], 'input_tokens': 0, 'output_tokens': 0, 'model': 'datalab-ocr'}
            raise Exception("No polling URL or markdown returned by Datalab.")
            
        for _ in range(60):
            time.sleep(5)
            poll_req = requests.get(request_check_url, headers=headers)
            poll_result = poll_req.json()
            if poll_result.get('status') == 'complete':
                return {'content': poll_result.get('markdown', ''), 'input_tokens': 0, 'output_tokens': 0, 'model': 'datalab-ocr'}
            elif poll_result.get('status') == 'error':
                raise Exception(f"Datalab Processing Error: {poll_result.get('error')}")
                
        raise Exception("Datalab OCR extraction timed out.")


class DocumentConverter:
    """Handles the conversion of documents (PDF, DOCX, etc.) to Markdown."""

    def __init__(self):
        try:
            import fitz
            self.fitz = fitz
        except ImportError:
            raise ImportError("PyMuPDF is required. Install with: pip install PyMuPDF")

    def count_tokens(self, text, model="gpt-4o"):
        """Count tokens in text using tiktoken."""
        if not HAS_TIKTOKEN:
            return len(text) // 4  # Rough estimate
        try:
            tk = get_tiktoken()
            if tk is None:
                return len(text) // 4
            encoding = tk.encoding_for_model(model)
            return len(encoding.encode(text))
        except:
            tk = get_tiktoken()
            if tk is None:
                return len(text) // 4
            encoding = tk.get_encoding("cl100k_base")
            return len(encoding.encode(text))

    def extract_metadata(self, pdf_path):
        """Extract PDF metadata."""
        doc = self.fitz.open(str(pdf_path))
        metadata = doc.metadata
        page_count = len(doc)
        doc.close()

        return {
            'title': metadata.get('title', ''),
            'author': metadata.get('author', ''),
            'subject': metadata.get('subject', ''),
            'creator': metadata.get('creator', ''),
            'creation_date': metadata.get('creationDate', ''),
            'modification_date': metadata.get('modDate', ''),
            'page_count': page_count
        }

    def extract_tables_pdfplumber(self, pdf_path, progress_callback=None, check_cancel=None):
        """Extract tables using pdfplumber with optimized parameters (2025 best practices).

        Uses tuned table detection settings for higher accuracy:
        - snap_x_tolerance: Improved column boundary detection
        - snap_y_tolerance: Better row alignment
        - join_x_tolerance/join_y_tolerance: Handle cell spanning
        - intersection_x/y_tolerance: Better grid detection
        """
        if not HAS_PDFPLUMBER:
            return {}

        # Optimized table detection settings (based on 2025 research showing 96% accuracy)
        table_settings = {
            "vertical_strategy": "lines_strict",  # Prioritize explicit lines
            "horizontal_strategy": "lines_strict",
            "snap_x_tolerance": 5,  # Key parameter for column alignment
            "snap_y_tolerance": 5,
            "join_x_tolerance": 5,
            "join_y_tolerance": 5,
            "edge_min_length": 10,
            "min_words_vertical": 1,
            "min_words_horizontal": 1,
            "intersection_x_tolerance": 5,
            "intersection_y_tolerance": 5,
        }

        # Fallback settings for tables without explicit borders
        fallback_settings = {
            "vertical_strategy": "text",
            "horizontal_strategy": "text",
            "snap_x_tolerance": 8,
            "snap_y_tolerance": 8,
            "join_x_tolerance": 8,
            "join_y_tolerance": 8,
            "min_words_vertical": 2,
            "min_words_horizontal": 1,
        }

        tables_by_page = {}
        try:
            _pdfplumber = get_pdfplumber()
            with _pdfplumber.open(str(pdf_path)) as pdf:
                total_pages = len(pdf.pages)
                for page_num, page in enumerate(pdf.pages):
                    if check_cancel and check_cancel():
                        return {}

                    if progress_callback:
                        progress_callback(page_num + 1, total_pages, f"Detecting tables (Page {page_num + 1}/{total_pages})...")

                    # Try strict line-based detection first
                    tables = page.extract_tables(table_settings)

                    # If no tables found, try text-based detection
                    if not tables:
                        tables = page.extract_tables(fallback_settings)

                    if tables:
                        # Filter out empty or single-cell "tables"
                        valid_tables = []
                        for table in tables:
                            if table and len(table) > 1 and any(len(row) > 1 for row in table if row):
                                valid_tables.append(table)
                        if valid_tables:
                            tables_by_page[page_num] = valid_tables
        except Exception as e:
            print(f"Table extraction warning: {e}")
            pass
        return tables_by_page

    def extract_text_tika(self, pdf_path):
        """Extract text using Apache Tika."""
        if not HAS_TIKA:
            return None
        
        try:
            # parser.from_file handles the server startup
            parsed = parser.from_file(str(pdf_path))
            content = parsed.get('content', '')
            if content:
                # Tika adds a lot of newlines
                return re.sub(r'\n{3,}', '\n\n', content).strip()
        except Exception as e:
            print(f"Tika extraction error: {e}")
        return None

    def extract_tables_tabula(self, pdf_path, total_pages, progress_callback=None, check_cancel=None):
        """Extract tables using Tabula-py."""
        if not HAS_TABULA:
            return {}
        
        tables_by_page = {}
        try:
            for page_num in range(1, total_pages + 1):
                if check_cancel and check_cancel():
                    return {}
                    
                if progress_callback:
                    progress_callback(page_num, total_pages, f"Tabula: Scanning page {page_num}...")
                
                try:
                    # lattice=True is often better for grid tables, stream=True for whitespace
                    # We accept defaults or try lattice
                    _tabula = get_tabula()
                    tables = _tabula.read_pdf(str(pdf_path), pages=str(page_num), multiple_tables=True, silent=True)
                    
                    if tables:
                        # Convert DataFrames to list of lists for consistency with pdfplumber
                        page_tables = []
                        for df in tables:
                            # Replace NaNs with empty string
                            df = df.fillna('')
                            # Get header
                            header = df.columns.tolist()
                            # Get values
                            values = df.values.tolist()
                            # Combine
                            clean_table = [header] + values
                            page_tables.append(clean_table)
                        
                        if page_tables:
                            tables_by_page[page_num - 1] = page_tables
                except Exception:
                    continue
                    
        except Exception as e:
            print(f"Tabula extraction warning: {e}")
            pass
        return tables_by_page

    def table_to_markdown(self, table):
        """Convert a table to markdown format."""
        if not table or len(table) < 1:
            return ""

        # Clean up cells
        cleaned = []
        for row in table:
            cleaned_row = [str(cell).replace('\n', ' ').strip() if cell else '' for cell in row]
            cleaned.append(cleaned_row)

        if not cleaned:
            return ""

        # Build markdown table
        md_lines = []

        # Header row
        header = cleaned[0]
        md_lines.append('| ' + ' | '.join(header) + ' |')
        md_lines.append('| ' + ' | '.join(['---'] * len(header)) + ' |')

        # Data rows
        for row in cleaned[1:]:
            # Pad row if necessary
            while len(row) < len(header):
                row.append('')
            md_lines.append('| ' + ' | '.join(row[:len(header)]) + ' |')

        return '\n'.join(md_lines)

    def ocr_page(self, page, dpi=300):
        """Perform OCR on a page image."""
        if not HAS_TESSERACT:
            return ""

        try:
            # Render page to image
            pix = page.get_pixmap(dpi=dpi)

            # Save to temp file
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as f:
                temp_path = f.name
                pix.save(temp_path)

            # OCR
            _Image = get_pil_image()
            _pytesseract = get_pytesseract()
            img = _Image.open(temp_path)
            text = _pytesseract.image_to_string(img)

            # Cleanup
            os.unlink(temp_path)

            return text
        except Exception as e:
            return f"[OCR Error: {str(e)}]"

    def page_to_base64(self, page, dpi=150):
        """Convert page to base64 encoded image."""
        pix = page.get_pixmap(dpi=dpi)
        img_bytes = pix.tobytes("png")
        return base64.b64encode(img_bytes).decode('utf-8')

    def extract_with_pymupdf4llm(self, pdf_path, output_path=None, extract_images=False,
                                  embed_images=False, image_dpi=150, image_format="png",
                                  page_chunks=False, progress_callback=None):
        """Extract PDF to Markdown using PyMuPDF4LLM (2025 best practice for LLM/RAG).

        PyMuPDF4LLM provides superior markdown extraction with:
        - Proper table detection and formatting
        - Header detection based on font size
        - Image extraction/embedding options
        - Reading order preservation
        - Page chunking for RAG applications

        Args:
            pdf_path: Path to PDF file
            output_path: Optional output path for images
            extract_images: Save images to disk
            embed_images: Embed images as base64 in markdown
            image_dpi: Resolution for image extraction (default 150)
            image_format: Image format (png, jpg)
            page_chunks: Return list of page chunks with metadata
            progress_callback: Progress callback function

        Returns:
            dict with 'markdown' (str or list), 'images' (list), 'metadata' (dict)
        """
        if not HAS_PYMUPDF4LLM:
            return None

        try:
            pdf_path = Path(pdf_path)
            images_dir = None

            if extract_images and output_path:
                images_dir = Path(output_path).parent / f"{Path(output_path).stem}_images"
                images_dir.mkdir(exist_ok=True)

            if progress_callback:
                progress_callback(0, 1, "Extracting with PyMuPDF4LLM...")

            # Build extraction options
            extract_options = {
                "page_chunks": page_chunks,
                "write_images": extract_images,
                "embed_images": embed_images,
                "dpi": image_dpi,
                "image_format": image_format,
            }

            if images_dir:
                extract_options["image_path"] = str(images_dir)

            # Extract markdown using PyMuPDF4LLM
            _pymupdf4llm = get_pymupdf4llm()
            result = _pymupdf4llm.to_markdown(str(pdf_path), **extract_options)

            if progress_callback:
                progress_callback(1, 1, "PyMuPDF4LLM extraction complete")

            # Handle page_chunks mode vs single string
            if page_chunks and isinstance(result, list):
                # Each chunk has: text, metadata (page, images, tables, etc.)
                markdown_parts = []
                all_images = []
                for chunk in result:
                    if isinstance(chunk, dict):
                        text = chunk.get('text', '')
                        page_num = chunk.get('metadata', {}).get('page', 0)
                        images = chunk.get('images', [])
                        # Add page marker for citation support
                        markdown_parts.append(f"[[PAGE_START: {page_num + 1}]]\n\n{text}")
                        all_images.extend(images)
                    else:
                        markdown_parts.append(str(chunk))

                return {
                    'markdown': "\n\n---\n\n".join(markdown_parts),
                    'chunks': result,
                    'images': all_images,
                    'method': 'pymupdf4llm'
                }
            else:
                return {
                    'markdown': result if isinstance(result, str) else str(result),
                    'chunks': None,
                    'images': [],
                    'method': 'pymupdf4llm'
                }

        except Exception as e:
            print(f"PyMuPDF4LLM extraction error: {e}")
            return None

    def extract_text_from_page(self, page):
        """Extract text from a PDF page with formatting."""
        blocks = page.get_text("dict")["blocks"]
        markdown_content = []

        for block in blocks:
            if block["type"] == 0:
                for line in block.get("lines", []):
                    line_text = ""
                    for span in line.get("spans", []):
                        text = span["text"]
                        flags = span["flags"]
                        is_bold = flags & 16
                        is_italic = flags & 2
                        if is_bold and is_italic:
                            text = f"***{text}***"
                        elif is_bold:
                            text = f"**{text}**"
                        elif is_italic:
                            text = f"*{text}*"
                        line_text += text
                    if line_text.strip():
                        markdown_content.append(line_text)

        return "\n".join(markdown_content)

    def detect_headers(self, page):
        """Detect potential headers based on font size analysis."""
        blocks = page.get_text("dict")["blocks"]
        font_sizes = []

        for block in blocks:
            if block["type"] == 0:
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        if span["text"].strip():
                            font_sizes.append(span["size"])

        if not font_sizes:
            return ""

        avg_size = sum(font_sizes) / len(font_sizes)
        max_size = max(font_sizes)

        markdown_lines = []
        for block in blocks:
            if block["type"] == 0:
                for line in block.get("lines", []):
                    line_text = ""
                    line_size = 0

                    for span in line.get("spans", []):
                        text = span["text"]
                        font_size = span["size"]
                        line_size = max(line_size, font_size)
                        flags = span["flags"]

                        is_bold = flags & 16
                        is_italic = flags & 2

                        if is_bold and is_italic:
                            text = f"***{text}***"
                        elif is_bold:
                            text = f"**{text}**"
                        elif is_italic:
                            text = f"*{text}*"

                        line_text += text

                    if line_text.strip():
                        if line_size >= max_size * 0.95:
                            line_text = f"# {line_text.strip('*')}"
                        elif line_size >= avg_size * 1.4:
                            line_text = f"## {line_text.strip('*')}"
                        elif line_size >= avg_size * 1.2:
                            line_text = f"### {line_text.strip('*')}"
                        markdown_lines.append(line_text)

        return "\n".join(markdown_lines)

    def is_scanned_pdf(self, pdf_path):
        """Check if PDF appears to be scanned (image-based)."""
        doc = self.fitz.open(str(pdf_path))
        text_chars = 0
        image_count = 0

        for page in doc:
            text_chars += len(page.get_text())
            image_count += len(page.get_images())

        doc.close()

        # If very little text but has images, likely scanned
        return text_chars < 100 and image_count > 0

    def convert_pdf_to_markdown(self, pdf_path, output_path=None, extract_images=False,
                                 detect_headers=True, progress_callback=None,
                                 use_ai=False, ai_client=None, ai_model=None,
                                 use_ocr=False, use_tables=True, use_vision=False,
                                 vision_model=None, custom_prompt=None,
                                 include_metadata=True, check_cancel=None,
                                 use_tika=False, use_tabula=False,
                                 use_pymupdf4llm=True, embed_images=False,
                                 checkpoint_dir=None):
        """Convert a PDF file to Markdown format with all features.

        2025 Best Practices Applied:
        - PyMuPDF4LLM for superior LLM/RAG-optimized extraction (default)
        - Enhanced table detection with tuned pdfplumber parameters
        - Image embedding option for self-contained markdown
        - Page markers for citation support
        """
        pdf_path = Path(pdf_path)
        cost_info = {'cost': 0, 'input_tokens': 0, 'output_tokens': 0}

        if not pdf_path.exists():
            raise FileNotFoundError(f"PDF file not found: {pdf_path}")

        # Validate file is a real PDF and not corrupt/password-protected
        try:
            _test_doc = self.fitz.open(str(pdf_path))
            if _test_doc.is_encrypted:
                _test_doc.close()
                raise PermissionError(
                    f"PDF is password-protected and cannot be opened: {pdf_path}. "
                    "Please provide an unprotected version of the file."
                )
            _page_count = len(_test_doc)
            _test_doc.close()
            
            # --- HIGH END CONVERSION OVERRIDE (DATALAB) ---
            if isinstance(ai_client, DatalabClient):
                if progress_callback:
                    progress_callback(0, 1, "Deploying DataLabs OCR High-End API...")
                result = ai_client.process_document(str(pdf_path))
                cost_info['cost'] += 0
                content = result['content']
                final_path = Path(output_path) if output_path else pdf_path.with_suffix('.md')
                with open(final_path, 'w', encoding='utf-8') as f:
                    f.write(content)
                return content, str(final_path), cost_info
            # ----------------------------------------------

            if _page_count == 0:
                raise ValueError(f"PDF has no pages: {pdf_path}")
        except PermissionError:
            raise
        except ValueError:
            raise
        except Exception as e:
            if "encrypted" in str(e).lower() or "password" in str(e).lower():
                raise PermissionError(
                    f"PDF is password-protected: {pdf_path}. "
                    "Please provide an unprotected version of the file."
                ) from e
            raise ValueError(
                f"Unable to open PDF (file may be corrupt): {pdf_path}. Error: {e}"
            ) from e

        if output_path is None:
            output_path = pdf_path.with_suffix('.md')
        else:
            output_path = Path(output_path)

        # Try PyMuPDF4LLM first (2025 best practice for LLM/RAG).
        # Skipped when the caller asked for OCR/vision/tika — those paths
        # need the per-page extraction loop further down.
        pymupdf4llm_result = None
        if (use_pymupdf4llm and HAS_PYMUPDF4LLM
                and not use_tika and not use_vision and not use_ocr):
            if progress_callback:
                progress_callback(0, 1, "Using PyMuPDF4LLM (optimized for LLM/RAG)...")

            pymupdf4llm_result = self.extract_with_pymupdf4llm(
                pdf_path,
                output_path=output_path,
                extract_images=extract_images,
                embed_images=embed_images,
                page_chunks=True,  # Enable for page markers
                progress_callback=progress_callback
            )

            # Reject empty extractions from scanned PDFs — pymupdf4llm
            # returns just page markers in that case, and feeding that to
            # the AI enhancer surfaces "[NO READABLE TEXT EXTRACTED]" as the
            # final output. Fall through to the traditional path so OCR /
            # vision can take over.
            if pymupdf4llm_result and pymupdf4llm_result.get('markdown'):
                _raw = pymupdf4llm_result['markdown']
                _stripped = re.sub(r'\[\[PAGE_START:[^\]]+\]\]', '', _raw)
                _stripped = re.sub(r'-{3,}', '', _stripped).strip()
                if len(_stripped) < 50 and self.is_scanned_pdf(pdf_path):
                    if progress_callback:
                        progress_callback(
                            0, 1,
                            "PyMuPDF4LLM extracted no text — scanned PDF detected, falling back to OCR/vision...")
                    pymupdf4llm_result = None

            if pymupdf4llm_result and pymupdf4llm_result.get('markdown'):
                # PyMuPDF4LLM successful - use its output
                raw_markdown = pymupdf4llm_result['markdown']

                # Extract metadata for header
                metadata = None
                if include_metadata:
                    metadata = self.extract_metadata(pdf_path)

                # Skip to AI enhancement if requested
                doc = self.fitz.open(str(pdf_path))
                total_pages = len(doc)
                doc.close()

                # Jump to AI enhancement section
                if use_ai and ai_client and ai_model:
                    if progress_callback:
                        progress_callback(total_pages, total_pages, "AI enhancing...")

                    def _swap_marker_for_hint(content, original_chunk):
                        # If AI judged the chunk unreadable, swap the
                        # literal marker for an actionable hint instead of
                        # writing "[NO READABLE TEXT EXTRACTED]" to the
                        # output file.
                        if "[NO READABLE TEXT EXTRACTED]" in content:
                            return ("[This page appears to be scanned/image-based. "
                                    "Enable Force OCR or switch to a vision-capable "
                                    "model to extract the text.]")
                        return content

                    try:
                        chunks = self._split_text(raw_markdown, max_chars=12000)
                        enhanced_parts, total_input_tokens, total_output_tokens = \
                            self._run_ai_chunks(
                                chunks, ai_client, ai_model, custom_prompt,
                                check_cancel=check_cancel,
                                progress_callback=progress_callback,
                                total_pages=total_pages,
                                checkpoint_dir=checkpoint_dir,
                                post_process=_swap_marker_for_hint)

                        markdown_content = "\n\n".join(enhanced_parts)
                        cost_info['input_tokens'] = total_input_tokens
                        cost_info['output_tokens'] = total_output_tokens
                    except Exception as e:
                        markdown_content = raw_markdown
                        raise Exception(f"AI enhancement failed: {str(e)}")
                else:
                    markdown_content = raw_markdown

                # Add metadata header
                if metadata and include_metadata:
                    meta_lines = ["---"]
                    if metadata.get('title'):
                        meta_lines.append(f"title: {metadata['title']}")
                    if metadata.get('author'):
                        meta_lines.append(f"author: {metadata['author']}")
                    if metadata.get('creation_date'):
                        meta_lines.append(f"date: {metadata['creation_date']}")
                    if metadata.get('page_count'):
                        meta_lines.append(f"pages: {metadata['page_count']}")
                    meta_lines.append(f"extraction_method: pymupdf4llm")
                    meta_lines.append("---\n")
                    markdown_content = "\n".join(meta_lines) + "\n" + markdown_content

                # Write output
                final_path = Path(output_path) if output_path else pdf_path.with_suffix('.md')
                with open(final_path, 'w', encoding='utf-8') as f:
                    f.write(markdown_content)

                return markdown_content, str(final_path), cost_info

        # Fallback to traditional extraction method
        doc = self.fitz.open(str(pdf_path))
        total_pages = len(doc)

        # Extract metadata
        metadata = None
        if include_metadata:
            metadata = self.extract_metadata(pdf_path)

        # Extract tables
        tables_by_page = {}
        if use_tables:
            if use_tabula and HAS_TABULA:
                if progress_callback:
                    progress_callback(0, total_pages, "Detecting tables with Tabula...")
                tables_by_page = self.extract_tables_tabula(pdf_path, total_pages, progress_callback, check_cancel)
            elif HAS_PDFPLUMBER:
                if progress_callback:
                    progress_callback(0, total_pages, "Detecting tables...")
                tables_by_page = self.extract_tables_pdfplumber(pdf_path, progress_callback, check_cancel)

        # Check if OCR is needed
        if use_ocr and self.is_scanned_pdf(pdf_path):
            if progress_callback:
                progress_callback(0, total_pages, "Scanned PDF detected, using OCR...")

        # Create images directory if extracting images
        images_dir = None
        if extract_images:
            images_dir = output_path.parent / f"{output_path.stem}_images"
            images_dir.mkdir(exist_ok=True)

        # Extract text from all pages
        all_text_parts = []
        total_input_tokens = 0
        total_output_tokens = 0

        # Tika Logic
        tika_content = None
        if use_tika and HAS_TIKA:
            if progress_callback:
                progress_callback(0, total_pages, "Extracting text with Tika...")
            tika_content = self.extract_text_tika(pdf_path)

        for page_num, page in enumerate(doc):
            if check_cancel and check_cancel():
                doc.close()
                raise Exception("Conversion cancelled")

            if progress_callback:
                progress_callback(page_num + 1, total_pages, "Extracting text...")

            page_text = ""
            if tika_content:
                pass  # Skip page-level extraction if using Tika
            
            # Use vision model for complex pages
            elif use_vision and ai_client and vision_model:
                try:
                    if progress_callback:
                        progress_callback(page_num + 1, total_pages, "Vision processing...")

                    img_base64 = self.page_to_base64(page)
                    result = ai_client.process_with_vision(img_base64, model=vision_model)
                    page_text = result['content']
                    total_input_tokens += result['input_tokens']
                    total_output_tokens += result['output_tokens']
                except Exception:
                    page_text = self.extract_text_from_page(page)

            # Use OCR for scanned pages
            elif use_ocr and HAS_TESSERACT:
                text = page.get_text()
                if len(text.strip()) < 50:  # Likely scanned
                    if progress_callback:
                        progress_callback(page_num + 1, total_pages, "OCR processing...")
                    page_text = self.ocr_page(page)
                else:
                    if detect_headers and not use_ai:
                        page_text = self.detect_headers(page)
                    else:
                        page_text = self.extract_text_from_page(page)
            else:
                if detect_headers and not use_ai:
                    page_text = self.detect_headers(page)
                else:
                    page_text = self.extract_text_from_page(page)

            # Insert tables from pdfplumber
            if page_num in tables_by_page:
                table_md = "\n\n".join([self.table_to_markdown(t) for t in tables_by_page[page_num]])
                if table_md:
                    page_text += f"\n\n{table_md}"

            # Prepend page tag for citation support
            page_content_with_tag = f"[[PAGE_START: {page_num + 1}]]\n\n{page_text}"
            all_text_parts.append(page_content_with_tag)

            # Extract images if requested
            if extract_images and images_dir:
                image_list = page.get_images()
                for img_idx, img in enumerate(image_list):
                    try:
                        xref = img[0]
                        base_image = doc.extract_image(xref)
                        image_bytes = base_image["image"]
                        image_ext = base_image["ext"]
                        image_filename = f"page{page_num + 1}_img{img_idx + 1}.{image_ext}"
                        image_path = images_dir / image_filename
                        with open(image_path, "wb") as img_file:
                            img_file.write(image_bytes)
                    except Exception:
                        pass

        doc.close()

        # Combine extracted text
        if tika_content:
            raw_markdown = tika_content
        else:
            raw_markdown = "\n\n---\n\n".join(all_text_parts)

        # Check if we have meaningful text content
        text_content = raw_markdown.replace("---", "").strip()
        has_meaningful_text = len(text_content) > 100  # At least 100 chars of real content

        # AI Enhancement
        if use_ai and ai_client and ai_model:
            if progress_callback:
                progress_callback(total_pages, total_pages, "AI enhancing...")

            try:
                # If no meaningful text and vision is available, suggest using vision mode
                if not has_meaningful_text:
                    if progress_callback:
                        progress_callback(total_pages, total_pages, "Low text detected, processing...")

                chunks = self._split_text(raw_markdown, max_chars=12000)

                def _handle_no_readable_marker(content, original_chunk):
                    if "[NO READABLE TEXT EXTRACTED]" in content:
                        if has_meaningful_text:
                            return original_chunk
                        return "[This page appears to be scanned/image-based. Enable OCR or Vision AI for better results.]"
                    return content

                enhanced_parts, chunk_input_tokens, chunk_output_tokens = \
                    self._run_ai_chunks(
                        chunks, ai_client, ai_model, custom_prompt,
                        check_cancel=check_cancel,
                        progress_callback=progress_callback,
                        total_pages=total_pages,
                        checkpoint_dir=checkpoint_dir,
                        post_process=_handle_no_readable_marker)
                total_input_tokens += chunk_input_tokens
                total_output_tokens += chunk_output_tokens

                markdown_content = "\n\n".join(enhanced_parts)

                # Final check - if AI produced garbage, use raw text
                if "Markdown Formatting Example" in markdown_content or "Table of Contents" in markdown_content[:500]:
                    if has_meaningful_text:
                        markdown_content = raw_markdown
                    else:
                        markdown_content = "[Document appears to be scanned/image-based. Enable OCR or Vision AI mode for text extraction.]"

            except Exception as e:
                markdown_content = raw_markdown
                raise Exception(f"AI enhancement failed: {str(e)}")
        else:
            markdown_content = raw_markdown

        # Add metadata header
        if metadata and include_metadata:
            meta_lines = ["---"]
            if metadata['title']:
                meta_lines.append(f"title: {metadata['title']}")
            if metadata['author']:
                meta_lines.append(f"author: {metadata['author']}")
            if metadata['creation_date']:
                meta_lines.append(f"date: {metadata['creation_date']}")
            meta_lines.append(f"pages: {metadata['page_count']}")
            meta_lines.append(f"source: {pdf_path.name}")
            meta_lines.append("---\n")
            markdown_content = "\n".join(meta_lines) + "\n" + markdown_content

        # Add image references
        if extract_images and images_dir:
            image_refs = []
            for img_file in sorted(images_dir.glob("*")):
                rel_path = f"{images_dir.name}/{img_file.name}"
                image_refs.append(f"![Image]({rel_path})")
            if image_refs:
                markdown_content += "\n\n## Images\n\n" + "\n\n".join(image_refs)

        # Clean up excessive newlines
        markdown_content = re.sub(r'\n{3,}', '\n\n', markdown_content)

        # Calculate cost
        cost = 0
        if ai_client and (total_input_tokens > 0 or total_output_tokens > 0):
            model_for_cost = vision_model if use_vision else ai_model
            cost = ai_client.calculate_cost(total_input_tokens, total_output_tokens, model_for_cost or ai_model)

        cost_info = {
            'cost': cost,
            'input_tokens': total_input_tokens,
            'output_tokens': total_output_tokens
        }

        # Save to file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(markdown_content)

        return markdown_content, output_path, cost_info

    def convert_docx_to_markdown(self, docx_path):
        """Convert DOCX to Markdown."""
        if not HAS_DOCX:
            return "Python-docx not installed.", {}

        try:
            doc = Document(docx_path)
            content = []
            
            # Extract core properties if possible
            props = doc.core_properties
            meta = []
            if props.title: meta.append(f"title: {props.title}")
            if props.author: meta.append(f"author: {props.author}")
            if props.created: meta.append(f"date: {props.created}")
            
            if meta:
                content.append("---\n" + "\n".join(meta) + "\n---\n")

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                # Simple style mapping
                style_name = para.style.name.lower()
                if 'title' in style_name:
                    content.append(f"# {text}")
                elif 'heading 1' in style_name:
                    content.append(f"# {text}")
                elif 'heading 2' in style_name:
                    content.append(f"## {text}")
                elif 'heading 3' in style_name:
                    content.append(f"### {text}")
                elif 'list bullet' in style_name:
                    content.append(f"* {text}")
                elif 'list number' in style_name:
                    content.append(f"1. {text}")
                else:
                    content.append(text)

            # Handle tables (naive approach: append them at the end or inline? 
            # DOCX structure is complex, this extracts tables separately or we need to iterate elements in order.
            # python-docx doesn't easily give document order of all elements mixed.
            # Simplified: just append text for now. Detailed reconstruction is complex.)
            
            return "\n\n".join(content), {'input_tokens': 0, 'output_tokens': 0}

        except Exception as e:
            raise Exception(f"DOCX conversion failed: {str(e)}")

    def convert_txt_to_markdown(self, txt_path):
        """Convert TXT to Markdown."""
        try:
            with open(txt_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            return text, {'input_tokens': 0, 'output_tokens': 0}
        except Exception as e:
            raise Exception(f"TXT conversion failed: {str(e)}")

    def convert_msg_to_markdown(self, msg_path):
        """Convert Outlook MSG to Markdown."""
        if not HAS_MSG:
            raise Exception("extract-msg not installed.")
            
        try:
            msg = extract_msg.Message(msg_path)
            content = []
            
            # Header
            content.append(f"# Subject: {msg.subject}")
            content.append(f"**From:** {msg.sender}")
            content.append(f"**To:** {msg.to}")
            if msg.date:
                content.append(f"**Date:** {msg.date}")
            content.append("\n---")
            
            # Body
            body = msg.body
            content.append(body)
            
            msg.close()
            return "\n\n".join(content), {'input_tokens': 0, 'output_tokens': 0}
        except Exception as e:
            raise Exception(f"MSG conversion failed: {str(e)}")

    def convert_eml_to_markdown(self, eml_path):
        """Convert EML to Markdown."""
        try:
            with open(eml_path, 'rb') as f:
                msg = BytesParser(policy=policy.default).parse(f)
            
            content = []
            content.append(f"# Subject: {msg['subject']}")
            content.append(f"**From:** {msg['from']}")
            content.append(f"**To:** {msg['to']}")
            if msg['date']:
                content.append(f"**Date:** {msg['date']}")
            content.append("\n---")
            
            body = msg.get_body(preferencelist=('plain', 'html'))
            if body:
                content.append(body.get_content())
                
            return "\n\n".join(content), {'input_tokens': 0, 'output_tokens': 0}
        except Exception as e:
            raise Exception(f"EML conversion failed: {str(e)}")

    def convert_epub_to_markdown(self, epub_path, use_ai=False, ai_client=None,
                                    ai_model=None, custom_prompt=None, progress_callback=None, **kwargs):
        """
        Convert EPUB to Markdown with enhanced processing (2025 best practices).

        Features:
        - BeautifulSoup for robust HTML parsing
        - Complete metadata extraction
        - Table of contents preservation
        - Table detection and conversion
        - Image reference handling
        - Footnote/endnote processing
        - Full content retention verification
        - AI enhancement support
        """
        if not HAS_EPUB:
            raise Exception("ebooklib not installed.")

        try:
            import html
            import re
            from bs4 import BeautifulSoup
            HAS_BS4 = True
        except ImportError:
            HAS_BS4 = False

        try:
            # Import epub from ebooklib (lazy loaded)
            import ebooklib
            from ebooklib import epub as epub_module
            book = epub_module.read_epub(epub_path)
            content = []
            toc_content = []
            footnotes = {}
            image_refs = []

            # ================================================================
            # ENHANCED METADATA EXTRACTION
            # ================================================================
            metadata_section = []

            # Title
            title = book.get_metadata('DC', 'title')
            if title:
                metadata_section.append(f"# {title[0][0]}")

            # Author(s)
            authors = book.get_metadata('DC', 'creator')
            if authors:
                author_names = [a[0] for a in authors]
                metadata_section.append(f"**Author(s):** {', '.join(author_names)}")

            # Publisher
            publisher = book.get_metadata('DC', 'publisher')
            if publisher:
                metadata_section.append(f"**Publisher:** {publisher[0][0]}")

            # Publication date
            date = book.get_metadata('DC', 'date')
            if date:
                metadata_section.append(f"**Date:** {date[0][0]}")

            # ISBN/Identifier
            identifier = book.get_metadata('DC', 'identifier')
            if identifier:
                metadata_section.append(f"**Identifier:** {identifier[0][0]}")

            # Language
            language = book.get_metadata('DC', 'language')
            if language:
                metadata_section.append(f"**Language:** {language[0][0]}")

            # Subject/Category
            subjects = book.get_metadata('DC', 'subject')
            if subjects:
                subject_list = [s[0] for s in subjects]
                metadata_section.append(f"**Subjects:** {', '.join(subject_list)}")

            # Description
            description = book.get_metadata('DC', 'description')
            if description:
                metadata_section.append(f"\n**Description:** {description[0][0]}")

            if metadata_section:
                content.append('\n'.join(metadata_section))
                content.append("\n---\n")

            # ================================================================
            # TABLE OF CONTENTS EXTRACTION
            # ================================================================
            try:
                toc = book.toc
                if toc:
                    toc_content.append("## Table of Contents\n")
                    for item in toc:
                        if isinstance(item, tuple):
                            # It's a section with sub-items
                            section, children = item
                            toc_content.append(f"- **{section.title}**")
                            for child in children:
                                toc_content.append(f"  - {child.title}")
                        else:
                            # Simple item
                            toc_content.append(f"- {item.title}")
                    toc_content.append("\n---\n")
                    content.append('\n'.join(toc_content))
            except Exception:
                pass  # TOC extraction failed, continue without it

            # ================================================================
            # CONTENT EXTRACTION WITH ENHANCED PARSING
            # ================================================================
            items_found = 0
            chapter_num = 0

            # Get items in spine order for correct sequence
            spine_items = []
            for item_id, linear in book.spine:
                item = book.get_item_with_id(item_id)
                if item:
                    spine_items.append(item)

            # Fallback to all items if spine is empty
            if not spine_items:
                spine_items = list(book.get_items_of_type(ebooklib.ITEM_DOCUMENT))

            for item in spine_items:
                if item.get_type() == ebooklib.ITEM_DOCUMENT:
                    items_found += 1
                    try:
                        # Get HTML content
                        try:
                            html_content = item.get_content().decode('utf-8')
                        except UnicodeDecodeError:
                            html_content = item.get_content().decode('utf-8', errors='ignore')

                        # Use BeautifulSoup if available (more robust)
                        if HAS_BS4:
                            text = self._parse_epub_html_bs4(html_content, footnotes, image_refs)
                        else:
                            text = self._parse_epub_html_regex(html_content)

                        if text and text.strip():
                            # Check if this looks like a chapter
                            if re.match(r'^#\s+', text) or len(text) > 500:
                                chapter_num += 1
                            content.append(text)

                    except Exception as item_error:
                        print(f"Warning: Failed to process EPUB item: {str(item_error)}")
                        continue

            # ================================================================
            # FOOTNOTES SECTION
            # ================================================================
            if footnotes:
                content.append("\n---\n## Notes\n")
                for ref_id, note_text in footnotes.items():
                    content.append(f"[^{ref_id}]: {note_text}")

            # Check if we found any content
            if items_found == 0:
                raise Exception("No document items found in EPUB file")

            if len(content) <= 2:  # Only metadata, no actual content
                raise Exception("No readable content extracted from EPUB file")

            raw_content = "\n\n".join(content)
            final_content = raw_content
            total_input_tokens = 0
            total_output_tokens = 0

            # AI Enhancement
            if use_ai and ai_client and ai_model:
                if progress_callback:
                    progress_callback(100, 100, "AI enhancing EPUB...")

                try:
                    chunks = self._split_text(raw_content, max_chars=12000)
                    enhanced_parts = []

                    for i, chunk in enumerate(chunks):
                        if progress_callback:
                            progress_callback(100, 100, f"AI chunk {i+1}/{len(chunks)}...")

                        result = ai_client.enhance_markdown(chunk, model=ai_model, custom_prompt=custom_prompt)
                        enhanced_parts.append(result['content'])
                        total_input_tokens += result['input_tokens']
                        total_output_tokens += result['output_tokens']

                    final_content = "\n\n".join(enhanced_parts)
                except Exception as e:
                    # Fallback to raw content on AI failure
                    final_content = raw_content
                    raise Exception(f"AI enhancement failed: {str(e)}")

            # Content retention check
            word_count = len(final_content.split())

            return final_content, {
                'input_tokens': total_input_tokens,
                'output_tokens': total_output_tokens,
                'chapters': chapter_num,
                'word_count': word_count,
                'has_toc': bool(toc_content),
                'footnotes_count': len(footnotes),
                'images_referenced': len(image_refs)
            }

        except Exception as e:
            raise Exception(f"EPUB conversion failed: {str(e)}")

    def _parse_epub_html_bs4(self, html_content: str, footnotes: dict, image_refs: list) -> str:
        """Parse EPUB HTML content using BeautifulSoup (robust parsing)."""
        from bs4 import BeautifulSoup
        import html as html_module

        soup = BeautifulSoup(html_content, 'html.parser')

        # Remove unwanted elements
        for tag in soup(['script', 'style', 'head', 'meta', 'link']):
            tag.decompose()

        # Extract and store footnotes/endnotes
        for note in soup.find_all(['aside', 'div'], class_=lambda x: x and ('note' in x.lower() or 'footnote' in x.lower())):
            note_id = note.get('id', '') or note.get('epub:type', '')
            if note_id:
                footnotes[note_id] = note.get_text(strip=True)
            note.decompose()

        # Handle images - extract references
        for img in soup.find_all('img'):
            src = img.get('src', '')
            alt = img.get('alt', 'Image')
            if src:
                image_refs.append({'src': src, 'alt': alt})
                img.replace_with(f"![{alt}]({src})")

        # Convert tables to markdown
        for table in soup.find_all('table'):
            md_table = self._table_to_markdown(table)
            table.replace_with(BeautifulSoup(f"\n{md_table}\n", 'html.parser'))

        # Convert headings
        for i in range(1, 7):
            for h in soup.find_all(f'h{i}'):
                prefix = '#' * i
                h.replace_with(f"\n{prefix} {h.get_text(strip=True)}\n")

        # Convert formatting
        for tag in soup.find_all(['strong', 'b']):
            tag.replace_with(f"**{tag.get_text()}**")
        for tag in soup.find_all(['em', 'i']):
            tag.replace_with(f"*{tag.get_text()}*")
        for tag in soup.find_all('u'):
            tag.replace_with(f"__{tag.get_text()}__")
        for tag in soup.find_all('code'):
            tag.replace_with(f"`{tag.get_text()}`")

        # Convert links
        for a in soup.find_all('a'):
            href = a.get('href', '')
            text = a.get_text(strip=True)
            if href.startswith('#'):
                # Internal link - might be footnote reference
                a.replace_with(f"[^{href[1:]}]" if 'note' in href.lower() else text)
            elif href:
                a.replace_with(f"[{text}]({href})")
            else:
                a.replace_with(text)

        # Convert lists
        for ul in soup.find_all('ul'):
            items = []
            for li in ul.find_all('li', recursive=False):
                items.append(f"- {li.get_text(strip=True)}")
            ul.replace_with('\n' + '\n'.join(items) + '\n')

        for ol in soup.find_all('ol'):
            items = []
            for idx, li in enumerate(ol.find_all('li', recursive=False), 1):
                items.append(f"{idx}. {li.get_text(strip=True)}")
            ol.replace_with('\n' + '\n'.join(items) + '\n')

        # Convert blockquotes
        for bq in soup.find_all('blockquote'):
            lines = bq.get_text().strip().split('\n')
            quoted = '\n'.join(f"> {line}" for line in lines)
            bq.replace_with(f"\n{quoted}\n")

        # Convert preformatted text
        for pre in soup.find_all('pre'):
            code = pre.get_text()
            pre.replace_with(f"\n```\n{code}\n```\n")

        # Convert horizontal rules
        for hr in soup.find_all('hr'):
            hr.replace_with("\n---\n")

        # Get text and clean up
        text = soup.get_text(separator='\n')
        text = html_module.unescape(text)

        # Clean up whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r'[ \t]+', ' ', text)
        text = re.sub(r' +\n', '\n', text)

        return text.strip()

    def _parse_epub_html_regex(self, html_content: str) -> str:
        """Parse EPUB HTML content using regex (fallback without BeautifulSoup)."""
        import html as html_module

        text = html_content

        # Remove scripts and styles
        text = re.sub(r'<script[^>]*>.*?</script>', '', text, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.DOTALL | re.IGNORECASE)

        # Convert headings
        for i in range(1, 7):
            prefix = '#' * i
            text = re.sub(rf'<h{i}[^>]*>(.*?)</h{i}>', rf'\n{prefix} \1\n', text, flags=re.IGNORECASE | re.DOTALL)

        # Convert formatting
        text = re.sub(r'<strong[^>]*>(.*?)</strong>', r'**\1**', text, flags=re.IGNORECASE | re.DOTALL)
        text = re.sub(r'<b[^>]*>(.*?)</b>', r'**\1**', text, flags=re.IGNORECASE | re.DOTALL)
        text = re.sub(r'<em[^>]*>(.*?)</em>', r'*\1*', text, flags=re.IGNORECASE | re.DOTALL)
        text = re.sub(r'<i[^>]*>(.*?)</i>', r'*\1*', text, flags=re.IGNORECASE | re.DOTALL)

        # Convert lists
        text = re.sub(r'<li[^>]*>(.*?)</li>', r'- \1\n', text, flags=re.IGNORECASE | re.DOTALL)
        text = re.sub(r'<[uo]l[^>]*>', '\n', text, flags=re.IGNORECASE)
        text = re.sub(r'</[uo]l>', '\n', text, flags=re.IGNORECASE)

        # Convert paragraphs and breaks
        text = re.sub(r'<p[^>]*>(.*?)</p>', r'\1\n\n', text, flags=re.IGNORECASE | re.DOTALL)
        text = re.sub(r'<br\s*/?>', '\n', text, flags=re.IGNORECASE)
        text = re.sub(r'<div[^>]*>', '\n', text, flags=re.IGNORECASE)
        text = re.sub(r'</div>', '\n', text, flags=re.IGNORECASE)
        text = re.sub(r'<hr\s*/?>', '\n---\n', text, flags=re.IGNORECASE)

        # Convert blockquotes
        text = re.sub(r'<blockquote[^>]*>(.*?)</blockquote>', r'\n> \1\n', text, flags=re.IGNORECASE | re.DOTALL)

        # Remove remaining tags
        text = re.sub(r'<[^>]+>', '', text)

        # Decode HTML entities
        text = html_module.unescape(text)

        # Clean up whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r'[ \t]+', ' ', text)

        return text.strip()

    def _table_to_markdown(self, table) -> str:
        """Convert HTML table to Markdown format."""
        rows = []
        headers = []

        # Extract headers
        thead = table.find('thead')
        if thead:
            for th in thead.find_all(['th', 'td']):
                headers.append(th.get_text(strip=True))

        # Extract body rows
        tbody = table.find('tbody') or table
        for tr in tbody.find_all('tr'):
            cells = []
            for td in tr.find_all(['td', 'th']):
                cells.append(td.get_text(strip=True))
            if cells:
                # First row might be headers if no thead
                if not headers and len(rows) == 0:
                    headers = cells
                else:
                    rows.append(cells)

        if not headers and not rows:
            return ""

        # Build markdown table
        md_lines = []

        # Header row
        if headers:
            md_lines.append("| " + " | ".join(headers) + " |")
            md_lines.append("| " + " | ".join(["---"] * len(headers)) + " |")

        # Data rows
        for row in rows:
            # Pad row if needed
            while len(row) < len(headers):
                row.append("")
            md_lines.append("| " + " | ".join(row[:len(headers)]) + " |")

        return '\n'.join(md_lines)

    def convert_mobi_to_markdown(self, mobi_path):
        """Convert MOBI to Markdown."""
        if not HAS_MOBI:
            raise Exception("mobi library not installed.")

        try:
            import html
            import re

            # Extract MOBI content
            tempdir, filepath = mobi.extract(mobi_path)

            # Read the extracted HTML
            content = []
            html_file = Path(tempdir) / filepath

            if html_file.exists():
                with open(html_file, 'r', encoding='utf-8', errors='ignore') as f:
                    html_content = f.read()

                # Enhanced HTML to markdown conversion (same as EPUB)
                text = html_content

                # Remove scripts and styles
                text = re.sub(r'<script[^>]*>.*?</script>', '', text, flags=re.DOTALL | re.IGNORECASE)
                text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.DOTALL | re.IGNORECASE)

                # Convert headings
                text = re.sub(r'<h1[^>]*>(.*?)</h1>', r'\n# \1\n', text, flags=re.IGNORECASE | re.DOTALL)
                text = re.sub(r'<h2[^>]*>(.*?)</h2>', r'\n## \1\n', text, flags=re.IGNORECASE | re.DOTALL)
                text = re.sub(r'<h3[^>]*>(.*?)</h3>', r'\n### \1\n', text, flags=re.IGNORECASE | re.DOTALL)
                text = re.sub(r'<h4[^>]*>(.*?)</h4>', r'\n#### \1\n', text, flags=re.IGNORECASE | re.DOTALL)
                text = re.sub(r'<h5[^>]*>(.*?)</h5>', r'\n##### \1\n', text, flags=re.IGNORECASE | re.DOTALL)
                text = re.sub(r'<h6[^>]*>(.*?)</h6>', r'\n###### \1\n', text, flags=re.IGNORECASE | re.DOTALL)

                # Convert formatting
                text = re.sub(r'<strong[^>]*>(.*?)</strong>', r'**\1**', text, flags=re.IGNORECASE | re.DOTALL)
                text = re.sub(r'<b[^>]*>(.*?)</b>', r'**\1**', text, flags=re.IGNORECASE | re.DOTALL)
                text = re.sub(r'<em[^>]*>(.*?)</em>', r'*\1*', text, flags=re.IGNORECASE | re.DOTALL)
                text = re.sub(r'<i[^>]*>(.*?)</i>', r'*\1*', text, flags=re.IGNORECASE | re.DOTALL)

                # Convert lists
                text = re.sub(r'<li[^>]*>(.*?)</li>', r'- \1\n', text, flags=re.IGNORECASE | re.DOTALL)
                text = re.sub(r'<ul[^>]*>', '\n', text, flags=re.IGNORECASE)
                text = re.sub(r'</ul>', '\n', text, flags=re.IGNORECASE)
                text = re.sub(r'<ol[^>]*>', '\n', text, flags=re.IGNORECASE)
                text = re.sub(r'</ol>', '\n', text, flags=re.IGNORECASE)

                # Convert paragraphs and breaks
                text = re.sub(r'<p[^>]*>(.*?)</p>', r'\1\n\n', text, flags=re.IGNORECASE | re.DOTALL)
                text = re.sub(r'<br\s*/?>', '\n', text, flags=re.IGNORECASE)
                text = re.sub(r'<div[^>]*>', '\n', text, flags=re.IGNORECASE)
                text = re.sub(r'</div>', '\n', text, flags=re.IGNORECASE)

                # Remove remaining tags
                text = re.sub(r'<[^>]+>', '', text)

                # Decode HTML entities
                text = html.unescape(text)

                # Clean up whitespace
                text = re.sub(r'\n{3,}', '\n\n', text)  # Normalize line breaks
                text = re.sub(r'[ \t]+', ' ', text)  # Normalize spaces
                text = text.strip()

                if text:
                    content.append(text)
            else:
                raise Exception("Failed to extract MOBI content - HTML file not found")

            if not content:
                raise Exception("No readable content extracted from MOBI file")

            return "\n\n".join(content), {'input_tokens': 0, 'output_tokens': 0}
        except Exception as e:
            raise Exception(f"MOBI conversion failed: {str(e)}")

    def convert_file(self, input_path, output_path=None, wrap_xml=False, **kwargs):
        """Dispatch conversion based on file extension."""
        path = Path(input_path)
        ext = path.suffix.lower()
        
        content = ""
        cost_info = {'cost': 0, 'input_tokens': 0, 'output_tokens': 0}
        
        if ext == '.pdf':
            # PDF conversion handles saving internally if output_path is provided
            # We intercept it to handle XML wrapping if needed
            if wrap_xml:
                 # Generate into memory first
                content, _, costs = self.convert_pdf_to_markdown(input_path, output_path=None, **kwargs)
                cost_info = costs
            else:
                return self.convert_pdf_to_markdown(input_path, output_path=output_path, **kwargs)
        
        elif ext == '.docx':
            content, cost_info = self.convert_docx_to_markdown(input_path)
        elif ext == '.txt':
            content, cost_info = self.convert_txt_to_markdown(input_path)
        elif ext == '.msg':
            content, cost_info = self.convert_msg_to_markdown(input_path)
        elif ext == '.eml':
            content, cost_info = self.convert_eml_to_markdown(input_path)
        elif ext == '.epub':
            content, cost_info = self.convert_epub_to_markdown(input_path, **kwargs)
        elif ext == '.mobi':
            content, cost_info = self.convert_mobi_to_markdown(input_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")
            
        # Wrap in XML if requested
        if wrap_xml:
            filename = path.name
            content = f'<document name="{filename}">\n{content}\n</document>'
            
        # Save to file if output path is provided
        final_path = output_path
        if output_path:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
        else:
            final_path = path.with_suffix('.md')

        # Update tokens in cost info to reflect actual content
        # Use fast estimation for large files to avoid GUI hanging
        if len(content) > 100000:
            # Fast estimation: ~4 chars per token
            cost_info['output_tokens'] = len(content) // 4
        else:
            # Accurate count for smaller files
            cost_info['output_tokens'] = self.count_tokens(content)

        return content, final_path, cost_info

    def _run_ai_chunks(self, chunks, ai_client, ai_model, custom_prompt=None,
                        check_cancel=None, progress_callback=None,
                        total_pages=1, checkpoint_dir=None, post_process=None):
        """Run AI enhancement over chunks with optional file-based checkpointing.

        When checkpoint_dir is provided, after each successful chunk we persist
        a single checkpoint.json (state + joined partial output + input hash)
        via atomic rename. On a subsequent run with the same checkpoint_dir
        and matching input hash, completed chunks are skipped and enhancement
        resumes from the last checkpoint. The atomic write guarantees disk
        state always reflects a consistent chunk boundary even across crashes.
        """
        import hashlib
        import os as _os

        resume_from = 0
        enhanced_parts = []
        input_tokens = 0
        output_tokens = 0

        # Hash the input chunks so re-uploading the same file with different
        # settings (e.g., OCR toggled) invalidates a stale checkpoint even
        # when the chunk count happens to match.
        input_hash = hashlib.sha256(
            "\0".join(chunks).encode('utf-8')).hexdigest()

        checkpoint_file = None
        if checkpoint_dir:
            checkpoint_dir = Path(checkpoint_dir)
            checkpoint_dir.mkdir(parents=True, exist_ok=True)
            checkpoint_file = checkpoint_dir / 'checkpoint.json'
            if checkpoint_file.exists():
                try:
                    state = json.loads(
                        checkpoint_file.read_text(encoding='utf-8'))
                    if (state.get('input_hash') == input_hash
                            and state.get('total_chunks') == len(chunks)):
                        resume_from = state.get('chunk_index', 0)
                        input_tokens = state.get('input_tokens', 0)
                        output_tokens = state.get('output_tokens', 0)
                        partial = state.get('partial', '')
                        if resume_from > 0 and partial:
                            enhanced_parts.append(partial)
                    else:
                        # Stale checkpoint — discard
                        checkpoint_file.unlink()
                except (json.JSONDecodeError, OSError):
                    resume_from = 0

        for i, chunk in enumerate(chunks):
            if i < resume_from:
                continue
            if check_cancel and check_cancel():
                raise Exception("Conversion cancelled")
            if progress_callback:
                progress_callback(total_pages, total_pages,
                                  f"AI chunk {i+1}/{len(chunks)}...")

            result = ai_client.enhance_markdown(
                chunk, model=ai_model, custom_prompt=custom_prompt)
            content = result['content']
            if post_process:
                content = post_process(content, chunk)
            enhanced_parts.append(content)
            input_tokens += result['input_tokens']
            output_tokens += result['output_tokens']

            if checkpoint_file:
                payload = json.dumps({
                    'input_hash': input_hash,
                    'chunk_index': i + 1,
                    'total_chunks': len(chunks),
                    'input_tokens': input_tokens,
                    'output_tokens': output_tokens,
                    'partial': "\n\n".join(enhanced_parts),
                })
                tmp = checkpoint_file.with_suffix('.json.tmp')
                tmp.write_text(payload, encoding='utf-8')
                _os.replace(tmp, checkpoint_file)

        return enhanced_parts, input_tokens, output_tokens

    def _split_text(self, text, max_chars=12000):
        """Split text into chunks for processing."""
        if len(text) <= max_chars:
            return [text]

        chunks = []
        current_chunk = ""
        pages = text.split("\n\n---\n\n")

        for page in pages:
            if len(current_chunk) + len(page) < max_chars:
                if current_chunk:
                    current_chunk += "\n\n---\n\n"
                current_chunk += page
            else:
                if current_chunk:
                    chunks.append(current_chunk)
                current_chunk = page

        if current_chunk:
            chunks.append(current_chunk)

        return chunks if chunks else [text]


class ExportManager:
    """Handles exporting to different formats."""

    @staticmethod
    def to_html(markdown_content, output_path):
        """Convert markdown to HTML."""
        try:
            import markdown as md
            html = md.markdown(markdown_content, extensions=['tables', 'fenced_code'])

            # Wrap in basic HTML document
            full_html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>Converted Document</title>
    <style>
        body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
               max-width: 800px; margin: 40px auto; padding: 0 20px; line-height: 1.6; }}
        table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
        th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
        th {{ background-color: #f5f5f5; }}
        code {{ background-color: #f5f5f5; padding: 2px 6px; border-radius: 3px; }}
        pre {{ background-color: #f5f5f5; padding: 15px; border-radius: 5px; overflow-x: auto; }}
        img {{ max-width: 100%; height: auto; }}
    </style>
</head>
<body>
{html}
</body>
</html>"""

            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(full_html)
            return True
        except Exception as e:
            raise Exception(f"HTML export failed: {str(e)}")

    @staticmethod
    def to_docx(markdown_content, output_path):
        """Convert markdown to DOCX."""
        if not HAS_DOCX:
            raise Exception("python-docx is required for DOCX export. Install with: pip install python-docx")

        try:
            doc = Document()

            lines = markdown_content.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue

                # Headers
                if line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('---'):
                    doc.add_paragraph('─' * 50)
                elif line.startswith('- ') or line.startswith('* '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif re.match(r'^\d+\. ', line):
                    doc.add_paragraph(re.sub(r'^\d+\. ', '', line), style='List Number')
                else:
                    # Handle bold and italic
                    p = doc.add_paragraph()
                    # Simple approach - just add text
                    clean_text = re.sub(r'\*{1,3}([^*]+)\*{1,3}', r'\1', line)
                    p.add_run(clean_text)

            doc.save(output_path)
            return True
        except Exception as e:
            raise Exception(f"DOCX export failed: {str(e)}")

    @staticmethod
    def to_txt(markdown_content, output_path):
        """Convert markdown to plain text."""
        try:
            # Remove markdown formatting
            text = markdown_content
            text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)  # Headers
            text = re.sub(r'\*{1,3}([^*]+)\*{1,3}', r'\1', text)  # Bold/italic
            text = re.sub(r'`([^`]+)`', r'\1', text)  # Inline code
            text = re.sub(r'!\[([^\]]*)\]\([^)]+\)', r'[Image: \1]', text)  # Images
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)  # Links
            text = re.sub(r'^[-*]\s+', '  * ', text, flags=re.MULTILINE)  # Lists
            text = re.sub(r'^\|.*\|$', '', text, flags=re.MULTILINE)  # Tables

            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)
            return True
        except Exception as e:
            raise Exception(f"TXT export failed: {str(e)}")


class MasterIndexGenerator:
    """Generates a master index for converted files."""
    
    @staticmethod
    def generate_index(folder_path, ai_client=None, model=None):
        folder = Path(folder_path)
        if not folder.exists():
            return "Folder not found."
            
        md_files = list(folder.glob("*.md"))
        if not md_files:
            return "No markdown files found."
            
        index_content = ["# Master Document Index\n"]
        total_tokens = 0
        
        # Sort files by name
        md_files.sort(key=lambda x: x.name)
        
        for md_file in md_files:
            if md_file.name == "master_index.md":
                continue
                
            try:
                with open(md_file, 'r', encoding='utf-8') as f:
                    content = f.read()
                    
                # Rough token count
                tokens = len(content) // 4
                if HAS_TIKTOKEN:
                    try:
                         tk = get_tiktoken()
                         if tk:
                             enc = tk.get_encoding("cl100k_base")
                             tokens = len(enc.encode(content))
                    except:
                        pass
                
                total_tokens += tokens
                
                summary = ""
                # Use AI for summary if available and cost effective (first 2k chars)
                if ai_client and model and len(content) > 100:
                    try:
                        prompt = f"Summarize the following document in one concise sentence:\n\n{content[:2000]}"
                        # Call API (simplified, assuming we can reuse client)
                         # Note: This is a synchronous call which might block UI if not careful.
                         # ideally this runs in a thread. 
                        pass 
                    except:
                        pass
                
                index_content.append(f"## [{md_file.name}]({md_file.name})")
                index_content.append(f"**Tokens:** {tokens}")
                # index_content.append(f"**Summary:** {summary}\n")
                index_content.append("")
                
            except Exception as e:
                 index_content.append(f"- {md_file.name} (Error reading file)")

        index_content.insert(1, f"**Total Documents:** {len(md_files)}")
        index_content.insert(2, f"**Total Project Tokens:** {total_tokens}\n---\n")

        output_file = folder / "master_index.md"
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write("\n".join(index_content))
            
        return output_file


class PDFToMarkdownApp:
    """Main application GUI with all features."""

    def __init__(self):
        # Use TkinterDnD if available for drag & drop support
        if HAS_DND:
            self.root = TkinterDnD.Tk()
        else:
            self.root = tk.Tk()

        self.root.title(f"Document to Markdown Converter Pro v{__version__}")
        self.root.geometry("950x850")
        self.root.minsize(850, 750)

        self.converter = None
        self.ai_client = None
        self.config = load_config()
        setup_java_environment(self.config)
        self.usage = load_usage()
        self.is_converting = False
        self.is_paused = False
        self.cancel_requested = False
        self.last_output_folder = None
        self.conversion_queue = queue.Queue()
        self.preview_content = ""

        # Handle window closing to ensure clean shutdown
        self.root.protocol("WM_DELETE_WINDOW", self.on_closing)

        self.setup_ui()
        self.update_usage_display()
        self.setup_drag_drop()

    def setup_drag_drop(self):
        """Setup drag and drop functionality."""
        if HAS_DND:
            self.file_tree.drop_target_register(DND_FILES)
            self.file_tree.dnd_bind('<<Drop>>', self.on_drop)

    def on_drop(self, event):
        """Handle dropped files."""
        files = self.root.tk.splitlist(event.data)
        for f in files:
            ext = Path(f).suffix.lower()
            if ext in ['.pdf', '.docx', '.txt', '.msg', '.eml', '.epub', '.mobi']:
                self.add_file_to_list(f)
        self.update_status()
        self.estimate_total_cost()

    def setup_ui(self):
        """Set up the user interface."""
        # Create notebook for tabs
        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(fill="both", expand=True, padx=5, pady=5)

        # Main conversion tab
        main_tab = ttk.Frame(self.notebook, padding="10")
        self.notebook.add(main_tab, text="Convert")

        # Settings tab
        settings_tab = ttk.Frame(self.notebook, padding="10")
        self.notebook.add(settings_tab, text="Settings")

        # Preview tab
        preview_tab = ttk.Frame(self.notebook, padding="10")
        self.notebook.add(preview_tab, text="Preview")

        # RAG/Vector tab
        rag_tab = ttk.Frame(self.notebook, padding="10")
        self.notebook.add(rag_tab, text="RAG/Vector")

        self.setup_main_tab(main_tab)
        self.setup_settings_tab(settings_tab)
        self.setup_preview_tab(preview_tab)
        self.setup_rag_tab(rag_tab)

    def setup_main_tab(self, parent):
        """Setup the main conversion tab."""
        parent.columnconfigure(0, weight=1)
        parent.rowconfigure(1, weight=1)

        # Title
        title_frame = ttk.Frame(parent)
        title_frame.grid(row=0, column=0, sticky="ew", pady=(0, 10))

        # Title
        title_frame = ttk.Frame(parent)
        title_frame.grid(row=0, column=0, sticky="ew", pady=(0, 10))

        title_label = ttk.Label(title_frame, text="Document to Markdown Converter Pro",
                                font=("Helvetica", 16, "bold"))
        title_label.pack(side="left")

        if HAS_DND:
            dnd_label = ttk.Label(title_frame, text="(Drag & Drop Enabled)", foreground="green")
            dnd_label.pack(side="left", padx=10)

            dnd_label.pack(side="left", padx=10)

        # File list frame
        file_frame = ttk.LabelFrame(parent, text="Files to Convert (drag files here)", padding="10")
        file_frame.grid(row=1, column=0, sticky="nsew", pady=5)
        file_frame.columnconfigure(0, weight=1)
        file_frame.rowconfigure(0, weight=1)

        # Treeview for file list
        columns = ("filename", "pages", "status", "cost")
        self.file_tree = ttk.Treeview(file_frame, columns=columns, show="headings", height=8)
        # Treeview for file list
        columns = ("filename", "pages", "status", "cost", "tokens")
        self.file_tree = ttk.Treeview(file_frame, columns=columns, show="headings", height=8)
        self.file_tree.heading("filename", text="File")
        self.file_tree.heading("pages", text="Pages")
        self.file_tree.heading("status", text="Status")
        self.file_tree.heading("cost", text="Est. Cost")
        self.file_tree.heading("tokens", text="Est. Tokens")
        self.file_tree.column("filename", width=300)
        self.file_tree.column("pages", width=50)
        self.file_tree.column("status", width=90)
        self.file_tree.column("cost", width=70)
        self.file_tree.column("tokens", width=70)

        # Context menu
        self.context_menu = tk.Menu(self.root, tearoff=0)
        self.context_menu.add_command(label="Remove", command=self.remove_selected)
        self.context_menu.add_command(label="Move Up", command=self.move_up)
        self.context_menu.add_command(label="Move Down", command=self.move_down)
        self.context_menu.add_separator()
        self.context_menu.add_command(label="Copy Text to Clipboard", command=self.copy_clipboard_selected)
        self.context_menu.add_command(label="Open PDF", command=self.open_selected_pdf)
        self.context_menu.add_command(label="Preview Output", command=self.preview_selected)

        self.file_tree.bind("<Button-3>", self.show_context_menu)

        tree_scroll = ttk.Scrollbar(file_frame, orient="vertical", command=self.file_tree.yview)
        self.file_tree.configure(yscrollcommand=tree_scroll.set)

        self.file_tree.grid(row=0, column=0, sticky="nsew")
        tree_scroll.grid(row=0, column=1, sticky="ns")

        # Button frame for file operations
        btn_frame = ttk.Frame(file_frame)
        btn_frame.grid(row=1, column=0, columnspan=2, pady=(10, 0))

        ttk.Button(btn_frame, text="Add Files...", command=self.add_files).pack(side="left", padx=3)
        ttk.Button(btn_frame, text="Add Folder...", command=self.add_folder).pack(side="left", padx=3)
        ttk.Button(btn_frame, text="Remove", command=self.remove_selected).pack(side="left", padx=3)
        ttk.Button(btn_frame, text="Clear All", command=self.clear_all).pack(side="left", padx=3)

        ttk.Separator(btn_frame, orient="vertical").pack(side="left", padx=10, fill="y")

        ttk.Button(btn_frame, text="Merge All", command=self.merge_all_pdfs).pack(side="left", padx=3)
        ttk.Button(btn_frame, text="Master Index", command=self.generate_master_index).pack(side="left", padx=3)

        # Quick options frame
        quick_frame = ttk.LabelFrame(parent, text="Quick Options", padding="10")
        quick_frame.grid(row=2, column=0, sticky="ew", pady=5)

        self.use_ai = tk.BooleanVar(value=self.config.get('use_ai_enhancement', True))
        ttk.Checkbutton(quick_frame, text="AI Enhancement", variable=self.use_ai,
                        command=self.on_option_change).pack(side="left", padx=10)

        self.use_ocr = tk.BooleanVar(value=self.config.get('use_ocr', False))
        ocr_state = "normal" if HAS_TESSERACT else "disabled"
        ocr_check = ttk.Checkbutton(quick_frame, text="OCR (scanned)", variable=self.use_ocr,
                                    command=self.on_option_change, state=ocr_state)
        ocr_check.pack(side="left", padx=10)

        self.use_tables = tk.BooleanVar(value=self.config.get('use_table_detection', True))
        table_state = "normal" if HAS_PDFPLUMBER else "disabled"
        ttk.Checkbutton(quick_frame, text="Table Detection", variable=self.use_tables,
                        state=table_state).pack(side="left", padx=10)

        self.use_vision = tk.BooleanVar(value=False)
        ttk.Checkbutton(quick_frame, text="Vision AI", variable=self.use_vision,
                        command=self.on_option_change).pack(side="left", padx=10)

        self.skip_existing = tk.BooleanVar(value=self.config.get('skip_existing', False))
        ttk.Checkbutton(quick_frame, text="Skip Existing", variable=self.skip_existing).pack(side="left", padx=10)

        # Output format selection
        ttk.Label(quick_frame, text="Format:").pack(side="left", padx=(20, 5))
        self.output_format = tk.StringVar(value=self.config.get('output_format', 'markdown'))
        format_combo = ttk.Combobox(quick_frame, textvariable=self.output_format,
                                    values=["markdown", "html", "docx", "txt"], width=10, state="readonly")
        format_combo.pack(side="left", padx=5)

        # Model selection frame
        model_frame = ttk.LabelFrame(parent, text="AI Settings", padding="10")
        model_frame.grid(row=3, column=0, sticky="ew", pady=5)
        model_frame.columnconfigure(1, weight=1)

        ttk.Label(model_frame, text="Model:").grid(row=0, column=0, sticky="w", padx=5)
        self.model_var = tk.StringVar(value=self.config.get('default_model', 'anthropic/claude-sonnet-4'))
        model_combo = ttk.Combobox(model_frame, textvariable=self.model_var, state="readonly", width=45)
        model_combo['values'] = [f"{k} - {v['name']}" for k, v in MODELS.items()]
        model_combo.grid(row=0, column=1, sticky="w", padx=5)
        model_combo.bind('<<ComboboxSelected>>', lambda e: self.estimate_total_cost())

        for k, v in MODELS.items():
            if k == self.model_var.get():
                model_combo.set(f"{k} - {v['name']}")
                break

        # Budget display
        ttk.Label(model_frame, text="Budget:").grid(row=0, column=2, sticky="w", padx=(20, 5))
        self.budget_label = ttk.Label(model_frame, text="$0.00 / $30.00")
        self.budget_label.grid(row=0, column=3, sticky="w")

        self.budget_progress = ttk.Progressbar(model_frame, length=150, mode='determinate')
        self.budget_progress.grid(row=0, column=4, padx=10)

        # Estimated cost
        ttk.Label(model_frame, text="Est. Total:").grid(row=1, column=0, sticky="w", padx=5, pady=(5, 0))
        self.est_cost_label = ttk.Label(model_frame, text="$0.0000", foreground="blue")
        self.est_cost_label.grid(row=1, column=1, sticky="w", padx=5, pady=(5, 0))

        # Output settings frame
        output_frame = ttk.LabelFrame(parent, text="Output Settings", padding="10")
        output_frame.grid(row=4, column=0, sticky="ew", pady=5)
        output_frame.columnconfigure(1, weight=1)

        self.output_same_dir = tk.BooleanVar(value=True)
        ttk.Radiobutton(output_frame, text="Same directory as PDF",
                        variable=self.output_same_dir, value=True,
                        command=self.toggle_output_dir).grid(row=0, column=0, sticky="w")
        
        self.wrap_xml = tk.BooleanVar(value=False)
        ttk.Checkbutton(output_frame, text="Wrap in XML Tags", variable=self.wrap_xml).grid(row=0, column=4, sticky="w", padx=10)

        ttk.Radiobutton(output_frame, text="Custom:",
                        variable=self.output_same_dir, value=False,
                        command=self.toggle_output_dir).grid(row=0, column=1, sticky="w")

        self.output_dir = tk.StringVar()
        self.output_dir_entry = ttk.Entry(output_frame, textvariable=self.output_dir, width=40, state="disabled")
        self.output_dir_entry.grid(row=0, column=2, sticky="ew", padx=5)

        self.browse_output_btn = ttk.Button(output_frame, text="Browse...",
                                            command=self.browse_output_dir, state="disabled")
        self.browse_output_btn.grid(row=0, column=3, padx=5)

        # Filename template
        ttk.Label(output_frame, text="Name Template:").grid(row=1, column=0, sticky="w", pady=(5, 0))
        self.name_template = tk.StringVar(value=self.config.get('output_template', '{name}'))
        template_entry = ttk.Entry(output_frame, textvariable=self.name_template, width=30)
        template_entry.grid(row=1, column=1, columnspan=2, sticky="w", padx=5, pady=(5, 0))
        ttk.Label(output_frame, text="({name}, {date}, {pages})", foreground="gray").grid(
            row=1, column=3, sticky="w", pady=(5, 0))

        # Progress frame
        progress_frame = ttk.Frame(parent)
        progress_frame.grid(row=5, column=0, sticky="ew", pady=10)
        progress_frame.columnconfigure(0, weight=1)

        ttk.Label(progress_frame, text="Progress:").grid(row=0, column=0, sticky="w")
        self.overall_progress_var = tk.DoubleVar()
        self.overall_progress_bar = ttk.Progressbar(progress_frame, variable=self.overall_progress_var,
                                                    maximum=100, mode='determinate')
        self.overall_progress_bar.grid(row=1, column=0, sticky="ew", pady=(2, 5))

        self.progress_label = ttk.Label(progress_frame, text="Ready")
        self.progress_label.grid(row=2, column=0)

        # Control buttons
        control_frame = ttk.Frame(parent)
        control_frame.grid(row=6, column=0, pady=10)

        self.convert_btn = ttk.Button(control_frame, text="Convert All",
                                      command=self.start_conversion)
        self.convert_btn.pack(side="left", padx=5)

        self.pause_btn = ttk.Button(control_frame, text="Pause", command=self.toggle_pause, state="disabled")
        self.pause_btn.pack(side="left", padx=5)

        self.cancel_btn = ttk.Button(control_frame, text="Cancel", command=self.cancel_conversion, state="disabled")
        self.cancel_btn.pack(side="left", padx=5)

        ttk.Separator(control_frame, orient="vertical").pack(side="left", padx=10, fill="y")

        self.open_folder_btn = ttk.Button(control_frame, text="Open Output Folder",
                                          command=self.open_output_folder, state="disabled")
        self.open_folder_btn.pack(side="left", padx=5)


        # Log frame
        log_frame = ttk.LabelFrame(parent, text="Log", padding="5")
        log_frame.grid(row=7, column=0, sticky="nsew", pady=5)
        log_frame.columnconfigure(0, weight=1)
        log_frame.rowconfigure(0, weight=1)
        parent.rowconfigure(7, weight=1)

        self.log_text = scrolledtext.ScrolledText(log_frame, wrap=tk.WORD, height=18)
        self.log_text.grid(row=0, column=0, sticky="nsew")

        # Status bar
        self.status_var = tk.StringVar(value="Add files to convert")
        status_bar = ttk.Label(parent, textvariable=self.status_var, relief="sunken")
        status_bar.grid(row=8, column=0, sticky="ew", pady=(5, 0))

    def safe_update(self, callback):
        """Thread-safe method to update GUI from worker threads."""
        try:
            self.root.after(0, callback)
            self.root.after(1, lambda: self.root.update_idletasks())
        except RuntimeError:
            # If main loop is not running, skip the update
            pass

    def setup_settings_tab(self, parent):
        """Setup the settings tab."""
        parent.columnconfigure(0, weight=1)

        # API Settings
        api_frame = ttk.LabelFrame(parent, text="API Configuration", padding="15")
        api_frame.grid(row=0, column=0, sticky="ew", pady=5)
        api_frame.columnconfigure(1, weight=1)

        ttk.Label(api_frame, text="OpenRouter API Key:").grid(row=0, column=0, sticky="w", pady=5)
        self.api_key_var = tk.StringVar(value=self.config.get('openrouter_api_key', ''))
        api_entry = ttk.Entry(api_frame, textvariable=self.api_key_var, width=60, show="*")
        api_entry.grid(row=0, column=1, sticky="ew", padx=5)

        ttk.Button(api_frame, text="Save", command=self.save_api_key).grid(row=0, column=2, padx=5)

        ttk.Label(api_frame, text="Monthly Budget ($):").grid(row=1, column=0, sticky="w", pady=5)
        self.budget_var = tk.StringVar(value=str(self.config.get('monthly_budget', 30.0)))
        budget_entry = ttk.Entry(api_frame, textvariable=self.budget_var, width=10)
        budget_entry.grid(row=1, column=1, sticky="w", padx=5)

        ttk.Button(api_frame, text="Reset Usage", command=self.reset_usage).grid(row=1, column=2, padx=5)

        # Custom Prompt
        prompt_frame = ttk.LabelFrame(parent, text="Custom AI Prompt", padding="15")
        prompt_frame.grid(row=1, column=0, sticky="nsew", pady=5)
        prompt_frame.columnconfigure(0, weight=1)
        prompt_frame.rowconfigure(0, weight=1)
        parent.rowconfigure(1, weight=1)

        self.custom_prompt_text = scrolledtext.ScrolledText(prompt_frame, wrap=tk.WORD, height=10)
        self.custom_prompt_text.grid(row=0, column=0, sticky="nsew")
        self.custom_prompt_text.insert("1.0", self.config.get('custom_prompt', '') or DEFAULT_PROMPT)

        prompt_btn_frame = ttk.Frame(prompt_frame)
        prompt_btn_frame.grid(row=1, column=0, pady=(10, 0))

        ttk.Button(prompt_btn_frame, text="Save Prompt", command=self.save_custom_prompt).pack(side="left", padx=5)
        ttk.Button(prompt_btn_frame, text="Reset to Default", command=self.reset_prompt).pack(side="left", padx=5)

        ttk.Label(prompt_frame, text="Use {text} as placeholder for PDF content",
                  foreground="gray").grid(row=2, column=0, pady=(5, 0))

        # Processing Settings
        proc_frame = ttk.LabelFrame(parent, text="Processing Settings", padding="15")
        proc_frame.grid(row=2, column=0, sticky="ew", pady=5)

        ttk.Label(proc_frame, text="Parallel Workers:").pack(side="left", padx=5)
        self.workers_var = tk.StringVar(value=str(self.config.get('parallel_workers', 2)))
        workers_spin = ttk.Spinbox(proc_frame, from_=1, to=8, width=5, textvariable=self.workers_var)
        workers_spin.pack(side="left", padx=5)

        self.include_metadata = tk.BooleanVar(value=self.config.get('include_metadata', True))
        ttk.Checkbutton(proc_frame, text="Include Metadata", variable=self.include_metadata).pack(side="left", padx=20)

        self.extract_images = tk.BooleanVar(value=True)
        ttk.Checkbutton(proc_frame, text="Extract Images", variable=self.extract_images).pack(side="left", padx=10)

        # Feature status
        status_frame = ttk.LabelFrame(parent, text="Feature Status", padding="15")
        status_frame.grid(row=3, column=0, sticky="ew", pady=5)

        features = [
            ("Drag & Drop", HAS_DND, "tkinterdnd2"),
            ("Table Detection", HAS_PDFPLUMBER, "pdfplumber"),
            ("OCR Support", HAS_TESSERACT, "pytesseract"),
            ("DOCX Export", HAS_DOCX, "python-docx"),
            ("DOCX/MSG Import", HAS_DOCX and HAS_MSG, "python-docx extract-msg"),
            ("Java", os.environ.get("JAVA_HOME") is not None, "JDK"),
            ("Tika (Better Text)", HAS_TIKA, "tika"),
            ("Tabula (Better Tables)", HAS_TABULA, "tabula-py")
        ]

        for i, (name, available, pkg) in enumerate(features):
            color = "green" if available else "red"
            status = "Available" if available else f"Install: pip install {pkg}"
            ttk.Label(status_frame, text=f"{name}:", font=("", 9, "bold")).grid(row=i, column=0, sticky="w", padx=5)
            ttk.Label(status_frame, text=status, foreground=color).grid(row=i, column=1, sticky="w", padx=5)

    def setup_preview_tab(self, parent):
        """Setup the preview tab."""
        parent.columnconfigure(0, weight=1)
        parent.rowconfigure(1, weight=1)

        # Preview controls
        ctrl_frame = ttk.Frame(parent)
        ctrl_frame.grid(row=0, column=0, sticky="ew", pady=5)

        ttk.Button(ctrl_frame, text="Preview Selected File", command=self.preview_selected).pack(side="left", padx=5)
        ttk.Button(ctrl_frame, text="Refresh", command=self.refresh_preview).pack(side="left", padx=5)

        # Preview text
        self.preview_text = scrolledtext.ScrolledText(parent, wrap=tk.WORD, font=("Consolas", 10))
        self.preview_text.grid(row=1, column=0, sticky="nsew")

    def setup_rag_tab(self, parent):
        """Setup the RAG/Vector tab for vector database management and search."""
        parent.columnconfigure(0, weight=1)
        parent.rowconfigure(2, weight=1)  # Results area gets extra space

        # Initialize RAG components
        self.vector_store = None
        self.embedding_client = None
        self.hybrid_retriever = None

        # ============== Status Frame ==============
        status_frame = ttk.LabelFrame(parent, text="Vector Database Status", padding="10")
        status_frame.grid(row=0, column=0, sticky="ew", pady=5)
        status_frame.columnconfigure(1, weight=1)

        # Status labels
        self.rag_status_var = tk.StringVar(value="Not initialized")
        self.rag_docs_var = tk.StringVar(value="Documents: 0")
        self.rag_chunks_var = tk.StringVar(value="Chunks: 0")
        self.rag_size_var = tk.StringVar(value="Size: 0 MB")

        ttk.Label(status_frame, text="Status:").grid(row=0, column=0, sticky="w", padx=5)
        ttk.Label(status_frame, textvariable=self.rag_status_var).grid(row=0, column=1, sticky="w", padx=5)
        ttk.Label(status_frame, textvariable=self.rag_docs_var).grid(row=0, column=2, sticky="w", padx=15)
        ttk.Label(status_frame, textvariable=self.rag_chunks_var).grid(row=0, column=3, sticky="w", padx=15)
        ttk.Label(status_frame, textvariable=self.rag_size_var).grid(row=0, column=4, sticky="w", padx=15)

        # Initialize button
        ttk.Button(status_frame, text="Initialize RAG", command=self.initialize_rag).grid(row=0, column=5, padx=10)
        ttk.Button(status_frame, text="Refresh Stats", command=self.refresh_rag_stats).grid(row=0, column=6, padx=5)

        # ============== Document Management Frame ==============
        doc_frame = ttk.LabelFrame(parent, text="Document Management", padding="10")
        doc_frame.grid(row=1, column=0, sticky="ew", pady=5)
        doc_frame.columnconfigure(0, weight=1)

        # Document list
        doc_list_frame = ttk.Frame(doc_frame)
        doc_list_frame.grid(row=0, column=0, sticky="ew")
        doc_list_frame.columnconfigure(0, weight=1)

        # Treeview for documents
        columns = ("source_file", "doc_type", "chunks", "date")
        self.rag_doc_tree = ttk.Treeview(doc_list_frame, columns=columns, show="headings", height=5)
        self.rag_doc_tree.heading("source_file", text="Document")
        self.rag_doc_tree.heading("doc_type", text="Type")
        self.rag_doc_tree.heading("chunks", text="Chunks")
        self.rag_doc_tree.heading("date", text="Date Added")
        self.rag_doc_tree.column("source_file", width=300)
        self.rag_doc_tree.column("doc_type", width=100)
        self.rag_doc_tree.column("chunks", width=80)
        self.rag_doc_tree.column("date", width=150)
        self.rag_doc_tree.grid(row=0, column=0, sticky="ew")

        # Scrollbar for doc list
        doc_scroll = ttk.Scrollbar(doc_list_frame, orient="vertical", command=self.rag_doc_tree.yview)
        doc_scroll.grid(row=0, column=1, sticky="ns")
        self.rag_doc_tree.configure(yscrollcommand=doc_scroll.set)

        # Document action buttons
        doc_btn_frame = ttk.Frame(doc_frame)
        doc_btn_frame.grid(row=1, column=0, sticky="w", pady=5)

        ttk.Button(doc_btn_frame, text="Vectorize Folder", command=self.vectorize_folder).pack(side="left", padx=5)
        ttk.Button(doc_btn_frame, text="Delete Selected", command=self.delete_from_vector_db).pack(side="left", padx=5)
        ttk.Button(doc_btn_frame, text="Refresh List", command=self.refresh_rag_doc_list).pack(side="left", padx=5)

        # Doc type selector
        ttk.Label(doc_btn_frame, text="Doc Type:").pack(side="left", padx=(20, 5))
        self.rag_doc_type_var = tk.StringVar(value=self.config.get('rag_settings', {}).get('default_doc_type', 'default'))
        doc_type_combo = ttk.Combobox(doc_btn_frame, textvariable=self.rag_doc_type_var, width=12, state="readonly")
        doc_type_combo['values'] = ['default', 'statute', 'case_law', 'treatise', 'rule', 'contract', 'pleading']
        doc_type_combo.pack(side="left", padx=5)

        # ============== Search Frame ==============
        search_frame = ttk.LabelFrame(parent, text="Retrieval Search", padding="10")
        search_frame.grid(row=2, column=0, sticky="nsew", pady=5)
        search_frame.columnconfigure(0, weight=1)
        search_frame.rowconfigure(2, weight=1)

        # Query input
        query_frame = ttk.Frame(search_frame)
        query_frame.grid(row=0, column=0, sticky="ew", pady=5)
        query_frame.columnconfigure(1, weight=1)

        ttk.Label(query_frame, text="Query:").grid(row=0, column=0, sticky="w", padx=5)
        self.rag_query_var = tk.StringVar()
        query_entry = ttk.Entry(query_frame, textvariable=self.rag_query_var, width=60)
        query_entry.grid(row=0, column=1, sticky="ew", padx=5)
        query_entry.bind('<Return>', lambda e: self.search_vector_db())

        ttk.Button(query_frame, text="Search", command=self.search_vector_db).grid(row=0, column=2, padx=5)

        # Search options
        opt_frame = ttk.Frame(search_frame)
        opt_frame.grid(row=1, column=0, sticky="w", pady=5)

        ttk.Label(opt_frame, text="k:").pack(side="left", padx=5)
        self.rag_k_var = tk.StringVar(value=str(self.config.get('rag_settings', {}).get('retrieval_k', 10)))
        k_spin = ttk.Spinbox(opt_frame, from_=1, to=50, width=5, textvariable=self.rag_k_var)
        k_spin.pack(side="left", padx=5)

        self.rag_hybrid_var = tk.BooleanVar(value=self.config.get('rag_settings', {}).get('hybrid_search', True))
        ttk.Checkbutton(opt_frame, text="Hybrid Search (BM25 + Dense)", variable=self.rag_hybrid_var).pack(side="left", padx=15)

        # Results treeview
        results_frame = ttk.Frame(search_frame)
        results_frame.grid(row=2, column=0, sticky="nsew", pady=5)
        results_frame.columnconfigure(0, weight=1)
        results_frame.rowconfigure(0, weight=1)

        result_cols = ("score", "source", "section", "preview")
        self.rag_results_tree = ttk.Treeview(results_frame, columns=result_cols, show="headings", height=8)
        self.rag_results_tree.heading("score", text="Score")
        self.rag_results_tree.heading("source", text="Source")
        self.rag_results_tree.heading("section", text="Section")
        self.rag_results_tree.heading("preview", text="Preview")
        self.rag_results_tree.column("score", width=60)
        self.rag_results_tree.column("source", width=150)
        self.rag_results_tree.column("section", width=150)
        self.rag_results_tree.column("preview", width=400)
        self.rag_results_tree.grid(row=0, column=0, sticky="nsew")
        self.rag_results_tree.bind('<<TreeviewSelect>>', self.on_rag_result_select)

        result_scroll = ttk.Scrollbar(results_frame, orient="vertical", command=self.rag_results_tree.yview)
        result_scroll.grid(row=0, column=1, sticky="ns")
        self.rag_results_tree.configure(yscrollcommand=result_scroll.set)

        # Chunk preview
        ttk.Label(search_frame, text="Full Chunk:").grid(row=3, column=0, sticky="w", pady=(10, 0))
        self.rag_chunk_preview = scrolledtext.ScrolledText(search_frame, wrap=tk.WORD, height=8, font=("Consolas", 9))
        self.rag_chunk_preview.grid(row=4, column=0, sticky="ew", pady=5)

        # Store search results for preview
        self.rag_search_results = []

        # ============== Settings Frame ==============
        settings_frame = ttk.LabelFrame(parent, text="RAG Settings", padding="10")
        settings_frame.grid(row=3, column=0, sticky="ew", pady=5)

        self.rag_auto_vectorize_var = tk.BooleanVar(value=self.config.get('rag_settings', {}).get('auto_vectorize', True))
        ttk.Checkbutton(settings_frame, text="Auto-vectorize after conversion",
                       variable=self.rag_auto_vectorize_var, command=self.save_rag_settings).pack(side="left", padx=10)

        self.rag_enabled_var = tk.BooleanVar(value=self.config.get('rag_settings', {}).get('enabled', True))
        ttk.Checkbutton(settings_frame, text="RAG Enabled",
                       variable=self.rag_enabled_var, command=self.save_rag_settings).pack(side="left", padx=10)

        # Check if RAG is available
        if not HAS_RAG:
            self.rag_status_var.set("RAG module not available - install dependencies")

    def initialize_rag(self):
        """Initialize the RAG components (vector store, embedding client)."""
        if not HAS_RAG:
            messagebox.showerror("Error", "RAG module not available.\n\nInstall dependencies:\npip install chromadb openai rank-bm25")
            return

        # Get OpenAI API key
        api_key = self.config.get('api_keys', {}).get('openai', '')
        if not api_key:
            # Try to use OpenRouter key if OpenAI not set
            api_key = self.config.get('api_keys', {}).get('openrouter', '')

        if not api_key:
            messagebox.showerror("Error", "No API key found.\n\nSet OpenAI or OpenRouter API key in Settings tab.")
            return

        try:
            self.rag_status_var.set("Initializing...")
            self.root.update()

            # Get vector DB path
            rag_settings = self.config.get('rag_settings', {})
            vector_db_path = rag_settings.get('vector_db_path', './vector_db')
            if not os.path.isabs(vector_db_path):
                vector_db_path = os.path.join(os.path.dirname(__file__), vector_db_path)

            # Initialize components (uses ChromaDB if available, falls back to numpy-based store)
            self.vector_store = get_vector_store(persist_dir=vector_db_path)
            self.embedding_client = EmbeddingClient(api_key=api_key, model=rag_settings.get('embedding_model', 'text-embedding-3-large'))
            self.hybrid_retriever = HybridRetriever(self.vector_store, self.embedding_client)

            self.rag_status_var.set("Ready")
            self.refresh_rag_stats()
            self.refresh_rag_doc_list()
            self.log("RAG system initialized successfully")

        except Exception as e:
            self.rag_status_var.set(f"Error: {str(e)[:50]}")
            messagebox.showerror("RAG Initialization Error", str(e))

    def refresh_rag_stats(self):
        """Refresh vector database statistics."""
        if not self.vector_store:
            return

        try:
            stats = self.vector_store.get_stats()
            self.rag_docs_var.set(f"Documents: {stats['total_documents']}")
            self.rag_chunks_var.set(f"Chunks: {stats['total_chunks']}")
            self.rag_size_var.set(f"Size: {stats['size_mb']:.1f} MB")
        except Exception as e:
            self.log(f"Error refreshing RAG stats: {e}")

    def refresh_rag_doc_list(self):
        """Refresh the document list in RAG tab."""
        if not self.vector_store:
            return

        # Clear existing
        for item in self.rag_doc_tree.get_children():
            self.rag_doc_tree.delete(item)

        try:
            docs = self.vector_store.list_documents()
            for doc in docs:
                self.rag_doc_tree.insert('', 'end', iid=doc['doc_id'], values=(
                    doc.get('source_file', 'Unknown'),
                    doc.get('doc_type', 'default'),
                    doc.get('chunk_count', 0),
                    doc.get('date', '')[:10] if doc.get('date') else ''
                ))
        except Exception as e:
            self.log(f"Error refreshing doc list: {e}")

    def vectorize_folder(self):
        """Vectorize all markdown files in a folder."""
        if not self.vector_store or not self.embedding_client:
            messagebox.showinfo("Info", "Please initialize RAG first (click 'Initialize RAG' button)")
            return

        folder = filedialog.askdirectory(title="Select folder with Markdown files")
        if not folder:
            return

        # Find all markdown files
        md_files = list(Path(folder).glob("*.md"))
        if not md_files:
            messagebox.showinfo("Info", "No markdown files found in folder")
            return

        doc_type = self.rag_doc_type_var.get()

        # Run vectorization in background
        def vectorize_all():
            total = len(md_files)
            success = 0
            for i, md_path in enumerate(md_files):
                try:
                    self.safe_update(lambda p=i, t=total: self.rag_status_var.set(f"Vectorizing {p+1}/{t}..."))
                    doc_id, chunks = vectorize_markdown_file(
                        str(md_path),
                        self.vector_store,
                        self.embedding_client,
                        doc_type=doc_type
                    )
                    if doc_id:
                        success += 1
                        self.safe_update(lambda: self.log(f"Vectorized: {md_path.name} ({chunks} chunks)"))
                except Exception as e:
                    self.safe_update(lambda: self.log(f"Error vectorizing {md_path.name}: {e}"))

            self.safe_update(lambda: self.rag_status_var.set("Ready"))
            self.safe_update(self.refresh_rag_stats)
            self.safe_update(self.refresh_rag_doc_list)
            self.safe_update(lambda: messagebox.showinfo("Complete", f"Vectorized {success}/{total} files"))

            # Refresh BM25 index for hybrid search
            if self.hybrid_retriever:
                self.hybrid_retriever.refresh_index()

        threading.Thread(target=vectorize_all, daemon=True).start()

    def delete_from_vector_db(self):
        """Delete selected documents from vector store."""
        if not self.vector_store:
            return

        selected = self.rag_doc_tree.selection()
        if not selected:
            messagebox.showinfo("Info", "Select documents to delete")
            return

        if not messagebox.askyesno("Confirm", f"Delete {len(selected)} document(s) from vector database?"):
            return

        for doc_id in selected:
            try:
                count = self.vector_store.delete_document(doc_id)
                self.log(f"Deleted document {doc_id} ({count} chunks)")
            except Exception as e:
                self.log(f"Error deleting {doc_id}: {e}")

        self.refresh_rag_stats()
        self.refresh_rag_doc_list()

        # Refresh BM25 index
        if self.hybrid_retriever:
            self.hybrid_retriever.refresh_index()

    def search_vector_db(self):
        """Execute a search query on the vector database."""
        if not self.hybrid_retriever:
            messagebox.showinfo("Info", "Please initialize RAG first")
            return

        query = self.rag_query_var.get().strip()
        if not query:
            return

        try:
            k = int(self.rag_k_var.get())
        except ValueError:
            k = 10

        use_hybrid = self.rag_hybrid_var.get()

        # Clear previous results
        for item in self.rag_results_tree.get_children():
            self.rag_results_tree.delete(item)
        self.rag_chunk_preview.delete('1.0', tk.END)

        try:
            self.rag_status_var.set("Searching...")
            self.root.update()

            results = self.hybrid_retriever.retrieve(query, k=k, use_hybrid=use_hybrid)
            self.rag_search_results = results

            for i, result in enumerate(results):
                # Parse section hierarchy from metadata
                section = ""
                if result.metadata.get('section_hierarchy'):
                    try:
                        hierarchy = json.loads(result.metadata['section_hierarchy']) if isinstance(result.metadata['section_hierarchy'], str) else result.metadata['section_hierarchy']
                        section = " > ".join(hierarchy[-2:]) if len(hierarchy) > 1 else (hierarchy[0] if hierarchy else "")
                    except:
                        section = ""

                # Truncate preview
                preview = result.chunk_text[:100].replace('\n', ' ') + "..."

                self.rag_results_tree.insert('', 'end', iid=str(i), values=(
                    f"{result.score:.3f}",
                    result.metadata.get('source_file', 'Unknown'),
                    section,
                    preview
                ))

            self.rag_status_var.set(f"Found {len(results)} results")
            self.log(f"Search '{query[:30]}...' returned {len(results)} results")

        except Exception as e:
            self.rag_status_var.set("Error")
            messagebox.showerror("Search Error", str(e))

    def on_rag_result_select(self, event):
        """Handle selection of a search result to show full chunk."""
        selected = self.rag_results_tree.selection()
        if not selected or not self.rag_search_results:
            return

        try:
            idx = int(selected[0])
            if 0 <= idx < len(self.rag_search_results):
                result = self.rag_search_results[idx]
                self.rag_chunk_preview.delete('1.0', tk.END)

                # Show metadata header
                meta_info = f"Source: {result.metadata.get('source_file', 'Unknown')}\n"
                meta_info += f"Score: {result.score:.4f} | Type: {result.metadata.get('doc_type', 'default')}\n"
                if result.metadata.get('page_range'):
                    meta_info += f"Pages: {result.metadata['page_range']}\n"
                meta_info += "-" * 50 + "\n\n"

                self.rag_chunk_preview.insert('1.0', meta_info + result.chunk_text)
        except Exception as e:
            pass

    def save_rag_settings(self):
        """Save RAG settings to config."""
        rag_settings = self.config.get('rag_settings', {})
        rag_settings['enabled'] = self.rag_enabled_var.get()
        rag_settings['auto_vectorize'] = self.rag_auto_vectorize_var.get()
        rag_settings['hybrid_search'] = self.rag_hybrid_var.get()
        rag_settings['default_doc_type'] = self.rag_doc_type_var.get()

        try:
            rag_settings['retrieval_k'] = int(self.rag_k_var.get())
        except ValueError:
            pass

        self.config['rag_settings'] = rag_settings
        save_config(self.config)

    def vectorize_document(self, md_path, metadata=None):
        """Vectorize a single converted markdown file (called after conversion)."""
        if not HAS_RAG or not self.vector_store or not self.embedding_client:
            return

        if not self.config.get('rag_settings', {}).get('auto_vectorize', True):
            return

        try:
            doc_type = self.config.get('rag_settings', {}).get('default_doc_type', 'default')
            doc_id, chunks = vectorize_markdown_file(
                str(md_path),
                self.vector_store,
                self.embedding_client,
                doc_type=doc_type,
                source_metadata=metadata
            )
            if doc_id:
                self.log(f"Auto-vectorized: {Path(md_path).name} ({chunks} chunks)")

                # Refresh index in background
                if self.hybrid_retriever:
                    threading.Thread(target=self.hybrid_retriever.refresh_index, daemon=True).start()

        except Exception as e:
            self.log(f"Auto-vectorization error: {e}")

    def show_context_menu(self, event):
        """Show context menu on right-click."""
        item = self.file_tree.identify_row(event.y)
        if item:
            self.file_tree.selection_set(item)
            self.context_menu.post(event.x_root, event.y_root)

    def move_up(self):
        """Move selected item up in the list."""
        selected = self.file_tree.selection()
        if not selected:
            return
        for item in selected:
            idx = self.file_tree.index(item)
            if idx > 0:
                self.file_tree.move(item, '', idx - 1)

    def move_down(self):
        """Move selected item down in the list."""
        selected = self.file_tree.selection()
        if not selected:
            return
        for item in reversed(selected):
            idx = self.file_tree.index(item)
            self.file_tree.move(item, '', idx + 1)

    def open_selected_pdf(self):
        """Open the selected PDF file."""
        selected = self.file_tree.selection()
        if selected:
            values = self.file_tree.item(selected[0])['values']
            if values:
                os.startfile(values[0])

    def preview_selected(self):
        """Preview the selected file's conversion."""
        selected = self.file_tree.selection()
        if not selected:
            messagebox.showinfo("No Selection", "Please select a file to preview.")
            return

        values = self.file_tree.item(selected[0])['values']
        if not values:
            return

        pdf_path = values[0]

        # Check if output already exists
        output_path = Path(self.get_output_path(pdf_path))
        if output_path.exists():
            with open(output_path, 'r', encoding='utf-8') as f:
                content = f.read()
            self.preview_text.delete("1.0", tk.END)
            self.preview_text.insert("1.0", content)
            self.notebook.select(2)  # Switch to preview tab
        else:
            messagebox.showinfo("Not Converted", "Convert the file first to see preview.")
            
    def copy_clipboard_selected(self):
        """Copy selected file conversion to clipboard."""
        selected = self.file_tree.selection()
        if not selected:
            return
            
        values = self.file_tree.item(selected[0])['values']
        pdf_path = values[0]
        output_path = Path(self.get_output_path(pdf_path))
        
        if output_path.exists():
            try:
                with open(output_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                self.root.clipboard_clear()
                self.root.clipboard_append(content)
                messagebox.showinfo("Copied", "Content copied to clipboard!")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to copy: {e}")
        else:
            messagebox.showinfo("Info", "File not converted yet.")
            
    def generate_master_index(self):
        """Generate master index for the output folder in a background thread."""
        if not self.last_output_folder:
             # Try to guess based on first item
             items = self.file_tree.get_children()
             if items:
                 val = self.file_tree.item(items[0])['values']
                 path = Path(val[0])
                 if self.output_same_dir.get():
                     self.last_output_folder = str(path.parent / "Markdown Converted Files")
        
        if not self.last_output_folder or not Path(self.last_output_folder).exists():
            messagebox.showinfo("Index", "No output folder available to index. Convert files first.")
            return

        def run_index():
            try:
                self.safe_update(lambda: self.log(f"Generating Master Index for: {self.last_output_folder}..."))
                outfile = MasterIndexGenerator.generate_index(self.last_output_folder)
                self.safe_update(lambda: self.log(f"Master Index generated: {outfile}"))
                self.safe_update(lambda: messagebox.showinfo("Success", f"Master Index generated at:\n{outfile}"))
                self.safe_update(lambda: os.startfile(outfile))
            except Exception as e:
                self.safe_update(lambda: self.log(f"Index generation failed: {e}"))
                self.safe_update(lambda: messagebox.showerror("Error", f"Failed to generate index: {e}"))

        thread = threading.Thread(target=run_index)
        thread.daemon = True
        thread.start()

    def refresh_preview(self):
        """Refresh the preview."""
        self.preview_selected()

    def merge_all_pdfs(self):
        """Merge all PDFs into a single markdown file."""
        items = self.file_tree.get_children()
        if len(items) < 2:
            messagebox.showinfo("Merge", "Add at least 2 PDF files to merge.")
            return

        output_file = filedialog.asksaveasfilename(
            title="Save Merged Markdown As",
            defaultextension=".md",
            filetypes=[("Markdown files", "*.md")]
        )

        if not output_file:
            return

        self.log("Starting merge operation...")

        try:
            if self.converter is None:
                self.converter = DocumentConverter()

            merged_content = []

            for idx, item in enumerate(items):
                values = self.file_tree.item(item)['values']
                pdf_path = values[0]

                self.log(f"Processing: {Path(pdf_path).name}")

                content, _, _ = self.converter.convert_pdf_to_markdown(
                    pdf_path,
                    output_path=None,
                    extract_images=False,
                    detect_headers=True,
                    use_ai=False
                )

                merged_content.append(f"# {Path(pdf_path).stem}\n\n{content}")

            final_content = "\n\n---\n\n".join(merged_content)

            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(final_content)

            self.log(f"Merged {len(items)} files to: {output_file}")
            self.last_output_folder = str(Path(output_file).parent)
            self.open_folder_btn.config(state="normal")

            messagebox.showinfo("Success", f"Merged {len(items)} PDFs successfully!")

        except Exception as e:
            self.log(f"Merge error: {str(e)}")
            messagebox.showerror("Error", f"Merge failed: {str(e)}")

    def on_option_change(self):
        """Handle option changes to update cost estimate."""
        self.estimate_total_cost()

    def estimate_total_cost(self):
        """Estimate total cost for all files."""
        if not self.use_ai.get():
            self.est_cost_label.config(text="$0.0000 (Local)")
            return

        items = self.file_tree.get_children()
        if not items:
            self.est_cost_label.config(text="$0.0000")
            return

        try:
            if self.converter is None:
                self.converter = DocumentConverter()

            total_chars = 0
            for item in items:
                values = self.file_tree.item(item)['values']
                pdf_path = values[0]

                # Estimate chars from page count (rough: 2000 chars per page)
                try:
                    import fitz
                    doc = fitz.open(pdf_path)
                    pages = len(doc)
                    doc.close()
                    total_chars += pages * 2000
                except:
                    total_chars += 10000  # Default estimate

            model = self.get_selected_model()
            if self.ai_client is None:
                provider = self.config.get('active_provider', 'openrouter')
                api_key = self.config.get('api_keys', {}).get(provider, '')
                if api_key:
                    self.ai_client = ClientFactory.get_client(provider, api_key)

            if self.ai_client:
                est_cost = self.ai_client.estimate_cost(total_chars, model)
                self.est_cost_label.config(text=f"${est_cost:.4f}")

                # Update individual file estimates
                for item in items:
                    values = self.file_tree.item(item)['values']
                    try:
                        import fitz
                        doc = fitz.open(values[0])
                        pages = len(doc)
                        doc.close()
                        file_cost = self.ai_client.estimate_cost(pages * 2000, model)
                        self.file_tree.item(item, values=(values[0], pages, values[2], f"~${file_cost:.4f}", values[4]))
                    except:
                        pass
            else:
                self.est_cost_label.config(text="No API key")

        except Exception as e:
            self.est_cost_label.config(text="Error")

    def save_api_key(self):
        """Save API settings to config."""
        # Check if we have the new multi-provider UI or old simple UI
        if hasattr(self, 'provider_var'):
            # New multi-provider UI
            provider = self.provider_var.get()
            key = self.api_key_var.get().strip()
            
            self.config['active_provider'] = provider
            if 'api_keys' not in self.config:
                self.config['api_keys'] = {}
            self.config['api_keys'][provider] = key
            
            # Update model list in main tab
            self.update_model_list()
            provider_name = PROVIDERS[provider]['name']
        else:
            # Old simple UI - just save as openrouter_api_key
            key = self.api_key_var.get().strip()
            self.config['openrouter_api_key'] = key
            # Ensure it's also saved to the new structure so clients pick it up
            if 'api_keys' not in self.config:
                self.config['api_keys'] = {}
            self.config['api_keys']['openrouter'] = key
            self.config['active_provider'] = 'openrouter'
            provider_name = "OpenRouter"
        
        try:
            self.config['monthly_budget'] = float(self.budget_var.get())
        except ValueError:
            pass
            
        save_config(self.config)
        self.ai_client = None  # Reset client
        
        self.log(f"Settings saved. Active Provider: {provider_name}")
        messagebox.showinfo("Saved", "Settings saved successfully!")

    def save_custom_prompt(self):
        """Save custom prompt to config."""
        prompt = self.custom_prompt_text.get("1.0", tk.END).strip()
        self.config['custom_prompt'] = prompt
        save_config(self.config)
        self.log("Custom prompt saved.")
        messagebox.showinfo("Saved", "Custom prompt saved!")

    def reset_prompt(self):
        """Reset prompt to default."""
        self.custom_prompt_text.delete("1.0", tk.END)
        self.custom_prompt_text.insert("1.0", DEFAULT_PROMPT)
        self.config['custom_prompt'] = ""
        save_config(self.config)

    def update_usage_display(self):
        """Update the budget/usage display."""
        self.usage = load_usage()
        budget = self.config.get('monthly_budget', 30.0)
        spent = self.usage.get('total_cost', 0)

        self.budget_label.config(text=f"${spent:.4f} / ${budget:.2f}")

        percentage = min((spent / budget) * 100, 100) if budget > 0 else 0
        self.budget_progress['value'] = percentage

        if percentage > 90:
            self.budget_label.config(foreground="red")
        elif percentage > 70:
            self.budget_label.config(foreground="orange")
        else:
            self.budget_label.config(foreground="green")

    def reset_usage(self):
        """Reset usage tracking."""
        if messagebox.askyesno("Reset Usage", "Reset monthly usage tracking to $0?"):
            self.usage = {
                "month": datetime.now().strftime("%Y-%m"),
                "total_cost": 0.0,
                "total_tokens": 0,
                "conversions": 0
            }
            save_usage(self.usage)
            self.update_usage_display()
            self.log("Usage tracking reset.")

    def track_usage(self, cost: float, tokens: int = 0):
        """Track usage for budget management.

        Args:
            cost: The cost incurred for this conversion
            tokens: The number of tokens used
        """
        # Ensure usage is current month
        current_month = datetime.now().strftime("%Y-%m")
        if self.usage.get("month") != current_month:
            self.usage = {
                "month": current_month,
                "total_cost": 0.0,
                "total_tokens": 0,
                "conversions": 0
            }

        # Update usage
        self.usage["total_cost"] = self.usage.get("total_cost", 0.0) + cost
        self.usage["total_tokens"] = self.usage.get("total_tokens", 0) + tokens
        self.usage["conversions"] = self.usage.get("conversions", 0) + 1

        # Save and update display
        save_usage(self.usage)
        self.update_usage_display()

        self.log(f"Usage tracked: ${cost:.4f} ({tokens} tokens) - Total: ${self.usage['total_cost']:.4f}")

    def update_provider_ui(self, event=None):
        """Update UI based on selected provider."""
        provider = self.provider_var.get()
        info = PROVIDERS.get(provider, PROVIDERS['openrouter'])
        
        self.api_key_label.config(text=f"{info['name']} Key:")
        
        # Load stored key
        stored_key = self.config.get('api_keys', {}).get(provider, "")
        self.api_key_var.set(stored_key)
        
        self.get_key_link.config(text="Get API Key")
        
    def open_provider_link(self):
        """Open the provider's API key page."""
        provider = self.provider_var.get()
        url = PROVIDERS.get(provider, {}).get('url', "")
        if url:
            import webbrowser
            webbrowser.open(url)
            
    def update_model_list(self):
        """Update the model combo box based on active provider."""
        provider = self.config.get('active_provider', 'openrouter')
        models = PROVIDER_MODELS.get(provider, {})
        
        # Update values
        if hasattr(self, 'model_combo'):
            self.model_combo['values'] = [f"{k} - {v['name']}" for k, v in models.items()]
            # Set default if current selection is invalid
            current = self.model_var.get()
            current_id = current.split(" - ")[0]
            if current_id not in models:
                 if models:
                    first = list(models.keys())[0]
                    self.model_var.set(f"{first} - {models[first]['name']}")
    
    def get_selected_model(self):
        """Get the selected model ID."""
        selection = self.model_var.get()
        provider = self.config.get('active_provider', 'openrouter')
        models = PROVIDER_MODELS.get(provider, {})

        # Extract the model part from selection (before " - ")
        model_part = selection.split(" - ")[0] if " - " in selection else selection

        # Direct match - model_part is already a valid model ID
        if model_part in models:
            return model_part

        # Check if selection starts with any valid model ID
        for model_id in models.keys():
            if selection.startswith(model_id):
                return model_id

        # Legacy model name handling - try to find a match by suffix
        # e.g., "gemini-2.5-flash" should match "google/gemini-2.5-flash"
        for model_id in models.keys():
            # Check if the model_part matches the end of any model_id (after the provider prefix)
            if "/" in model_id:
                _, model_name = model_id.split("/", 1)
                if model_part == model_name or model_part.startswith(model_name):
                    return model_id

        # If still no match and using openrouter, try to add common prefixes
        if provider == 'openrouter':
            prefixes = ['google/', 'anthropic/', 'openai/', 'deepseek/', 'meta-llama/', 'qwen/', 'mistralai/']
            for prefix in prefixes:
                potential_id = prefix + model_part
                if potential_id in models:
                    return potential_id

        # FINAL FALLBACK: Return first available model if nothing matches
        # This prevents sending invalid model IDs to the API
        if models:
            default_model = list(models.keys())[0]
            print(f"Warning: Model '{model_part}' not found, using default: {default_model}")
            return default_model

        return model_part

    def add_files(self):
        """Add multiple files to the list."""
        filenames = filedialog.askopenfilenames(
            title="Select Files",
            filetypes=[
                ("All Supported", "*.pdf;*.docx;*.txt;*.msg;*.eml;*.epub;*.mobi"),
                ("PDF files", "*.pdf"),
                ("Word Documents", "*.docx"),
                ("Text Files", "*.txt"),
                ("Email Messages", "*.msg;*.eml"),
                ("Ebook Files", "*.epub;*.mobi"),
                ("All files", "*.*")
            ]
        )
        for filename in filenames:
            self.add_file_to_list(filename)
        self.update_status()
        self.estimate_total_cost()

    def add_folder(self):
        """Add all supported files from a folder."""
        folder = filedialog.askdirectory(title="Select Folder")
        if folder:
            folder_path = Path(folder)
            extensions = ['*.pdf', '*.docx', '*.txt', '*.msg', '*.eml', '*.epub', '*.mobi']
            files = []
            for ext in extensions:
                files.extend(list(folder_path.glob(ext)) + list(folder_path.glob(f"**/{ext}")))
            
            for file_path in files:
                self.add_file_to_list(str(file_path))
            self.update_status()
            self.estimate_total_cost()
            self.update_status()
            self.estimate_total_cost()
            if not files:
                messagebox.showinfo("No Files Found", "No supported files were found in the selected folder.")

    def add_file_to_list(self, filepath):
        """Add a single file to the list if not already present."""
        for item in self.file_tree.get_children():
            values = self.file_tree.item(item)['values']
            if values and values[0] == filepath:
                return

        # Get page count
        try:
            import fitz
            doc = fitz.open(filepath)
            pages = len(doc)
            doc.close()
        except:
            pages = "?"

        self.file_tree.insert("", "end", values=(filepath, pages, "Pending", "-", "-"))

    def get_output_path(self, input_path):
        """Get the output path for a given input file."""
        input_path = Path(input_path)

        # Apply naming template
        template = self.name_template.get() or "{name}"
        name = template.replace("{name}", input_path.stem)
        name = name.replace("{date}", datetime.now().strftime("%Y%m%d"))

        try:
            import fitz
            doc = fitz.open(str(input_path))
            pages = len(doc)
            doc.close()
            name = name.replace("{pages}", str(pages))
        except:
            name = name.replace("{pages}", "0")

        # Get extension based on format
        fmt = self.output_format.get()
        ext = {"markdown": ".md", "html": ".html", "docx": ".docx", "txt": ".txt"}.get(fmt, ".md")

        if self.output_same_dir.get():
             # Create "Markdown Converted Files" folder
            output_dir = input_path.parent / "Markdown Converted Files"
            output_dir.mkdir(parents=True, exist_ok=True)
            return str(output_dir / (name + ext))
        else:
            output_dir = self.output_dir.get()
            if output_dir:
                return str(Path(output_dir) / (name + ext))
            
            # Default fallback if custom dir not set
            output_dir = input_path.parent / "Markdown Converted Files"
            output_dir.mkdir(parents=True, exist_ok=True)
            return str(output_dir / (name + ext))

    def toggle_output_dir(self):
        """Toggle the output directory entry state."""
        if self.output_same_dir.get():
            self.output_dir_entry.config(state="disabled")
            self.browse_output_btn.config(state="disabled")
        else:
            self.output_dir_entry.config(state="normal")
            self.browse_output_btn.config(state="normal")

    def browse_output_dir(self):
        """Browse for output directory."""
        folder = filedialog.askdirectory(title="Select Output Directory")
        if folder:
            self.output_dir.set(folder)

    def open_output_folder(self):
        """Open the output folder in Windows Explorer."""
        if self.last_output_folder and Path(self.last_output_folder).exists():
            os.startfile(self.last_output_folder)
        else:
            messagebox.showinfo("No Output", "No output folder available. Convert some files first.")

    def remove_selected(self):
        """Remove selected items from the list."""
        selected = self.file_tree.selection()
        for item in selected:
            self.file_tree.delete(item)
        self.update_status()
        self.estimate_total_cost()

    def clear_all(self):
        """Clear all items from the list."""
        for item in self.file_tree.get_children():
            self.file_tree.delete(item)
        self.update_status()
        self.estimate_total_cost()

    def update_status(self):
        """Update the status bar."""
        count = len(self.file_tree.get_children())
        if count == 0:
            self.status_var.set("Add files to convert (or drag & drop)")
        else:
            mode_parts = []
            if self.use_ai.get():
                mode_parts.append("AI")
            if self.use_ocr.get():
                mode_parts.append("OCR")
            if self.use_vision.get():
                mode_parts.append("Vision")
            mode = "+".join(mode_parts) if mode_parts else "Local"
            self.status_var.set(f"{count} file(s) ready ({mode} mode)")

    def log(self, message):
        """Add a message to the log."""
        timestamp = datetime.now().strftime("%H:%M:%S")
        self.log_text.insert(tk.END, f"[{timestamp}] {message}\n")
        self.log_text.see(tk.END)
        self.root.update_idletasks()

    def update_item_status(self, item, status, cost="-", tokens="-"):
        """Update the status of a tree item."""
        values = self.file_tree.item(item)['values']
        if values:
            self.file_tree.item(item, values=(values[0], values[1], status, cost, tokens))
        self.root.update_idletasks()

    def check_budget(self):
        """Check if budget allows for more conversions."""
        budget = self.config.get('monthly_budget', 30.0)
        spent = self.usage.get('total_cost', 0)
        return spent < budget

    def toggle_pause(self):
        """Toggle pause state."""
        self.is_paused = not self.is_paused
        self.pause_btn.config(text="Resume" if self.is_paused else "Pause")
        self.log("Paused" if self.is_paused else "Resumed")

    def cancel_conversion(self):
        """Cancel the conversion."""
        self.cancel_requested = True
        self.log("Cancellation requested...")

    def start_conversion(self):
        """Start the batch conversion process."""
        items = self.file_tree.get_children()
        if not items:
            messagebox.showerror("Error", "Please add files to convert.")
            return

        if self.is_converting:
            self.log("Already converting, ignoring click.")
            return

        if self.use_ai.get() and not self.check_budget():
            messagebox.showwarning("Budget Exceeded",
                "Monthly budget exceeded. Disable AI enhancement or reset usage.")
            return

        # Clear log and show immediate feedback
        self.log_text.delete(1.0, tk.END)
        self.log("="*60)
        self.log("STARTING BATCH CONVERSION")
        self.log("="*60)
        self.log(f"Files to convert: {len(items)}")

        self.is_converting = True
        self.is_paused = False
        self.cancel_requested = False
        self.convert_btn.config(state="disabled")
        self.pause_btn.config(state="normal")
        self.cancel_btn.config(state="normal")
        self.overall_progress_var.set(0)
        self.progress_label.config(text="Initializing...")
        self.root.update_idletasks()  # Force immediate GUI update

        for item in items:
            values = self.file_tree.item(item)['values']
            self.update_item_status(item, "Pending", values[3] if len(values) > 3 else "-", "-")

        self.log("Starting worker thread...")
        self.root.update_idletasks()  # Force update before thread starts

        thread = threading.Thread(target=self.run_batch_conversion)
        thread.daemon = True
        thread.start()

        self.log("Worker thread started - processing files...")
        self.root.update_idletasks()  # Force update after thread starts

    def run_batch_conversion(self):
        """Run the batch conversion."""
        try:
            self.safe_update(lambda: self.log("Starting batch conversion..."))
            
            # Reload config to ensure we have any keys saved just before clicking convert
            self.config = load_config()

            if self.converter is None:
                self.converter = DocumentConverter()

            use_ai = self.use_ai.get()
            ai_client = None
            ai_model = None

            if use_ai:
                provider = self.config.get('active_provider', 'openrouter')
                api_keys = self.config.get('api_keys', {})
                api_key = api_keys.get(provider, '')
                
                # Fallback
                if not api_key and provider == 'openrouter':
                    api_key = self.config.get('openrouter_api_key', '')
                
                # Debug log
                if api_key:
                    masked_key = f"{api_key[:10]}..." if len(api_key) > 10 else "***"
                    self.safe_update(lambda: self.log(f"Using Provider: {provider}, Key: {masked_key}"))
                else:
                    self.safe_update(lambda: self.log(f"Using Provider: {provider}, Key: [NOT FOUND]"))

                if api_key:
                    try:
                        ai_client = ClientFactory.get_client(provider, api_key)
                        ai_model = self.get_selected_model()
                        self.safe_update(lambda: self.log(f"Using AI model: {ai_model}"))
                    except Exception as e:
                        self.safe_update(lambda err=str(e): self.log(f"Failed to initialize AI client: {err}"))
                        self.safe_update(lambda: self.log("Falling back to local processing."))
                        use_ai = False
                else:
                    self.safe_update(lambda: self.log("No API key. Using local processing."))
                    use_ai = False

            items = self.file_tree.get_children()
            total_files = len(items)
            successful = 0
            failed = 0
            skipped = 0
            total_cost = 0

            custom_prompt = self.config.get('custom_prompt', '') or None
            vision_model = self.config.get('vision_model', 'openai/gpt-4o-mini')

            try:
                max_workers = int(self.workers_var.get())
            except ValueError:
                max_workers = 1
                
            self.safe_update(lambda: self.log(f"Starting batch conversion with {max_workers} workers..."))

            # We need to gather all data needed for processing BEFORE starting threads
            # because accessing Tkinter widgets (treeview items) from threads is unsafe.
            work_items = []
            for item in items:
                values = self.file_tree.item(item)['values']
                input_path = values[0]
                output_path = self.get_output_path(input_path)
                work_items.append({
                    'id': item,
                    'input_path': input_path,
                    'filename': Path(input_path).name,
                    'output_path': output_path,
                    'skip_existing': self.skip_existing.get() and Path(output_path).exists()
                })

            completed_count = 0
            
            with ThreadPoolExecutor(max_workers=max_workers) as executor:
                futures = {}
                
                for work_item in work_items:
                    if self.cancel_requested:
                        break
                        
                    # If file should be skipped, handle it in the main thread
                    if work_item['skip_existing']:
                        self.safe_update(lambda i=work_item['id']: self.update_item_status(i, "Skipped", "-", "-"))
                        skipped += 1
                        completed_count += 1
                        progress = (completed_count / total_files) * 100
                        self.safe_update(lambda p=progress: self.overall_progress_var.set(p))
                        continue

                    # Update last_output_folder for the first non-skipped item
                    if not self.last_output_folder:
                        self.last_output_folder = str(Path(work_item['output_path']).parent)

                    # Update status to converting for this item
                    self.safe_update(lambda i=work_item['id'], f=work_item['filename']: self.progress_label.config(text=f"Converting: {f}"))
                    self.safe_update(lambda i=work_item['id']: self.update_item_status(i, "Converting...", "-", "-"))

                    future = executor.submit(self.process_single_file, work_item, use_ai, ai_client, ai_model, vision_model, custom_prompt)
                    futures[future] = work_item

                for future in as_completed(futures):
                    if self.cancel_requested:
                        executor.shutdown(wait=False)
                        break
                        
                    work_item = futures[future]
                    try:
                        success, result_cost = future.result()
                        if success:
                            successful += 1
                            total_cost += result_cost
                        else:
                            failed += 1
                    except Exception as e:
                        failed += 1
                        self.safe_update(lambda e=e: self.log(f"Thread error: {e}"))
                    
                    completed_count += 1
                    progress = (completed_count / total_files) * 100
                    self.safe_update(lambda p=progress: self.overall_progress_var.set(p))
                    self.root.update_idletasks()

            self.safe_update(lambda s=successful, f=failed, sk=skipped, tc=total_cost:
                self.conversion_batch_complete(s, f, sk, tc))

        except Exception as e:
            self.safe_update(lambda err=str(e): self.log(f"Critical error: {err}"))
            total = len(self.file_tree.get_children())
            self.safe_update(lambda t=total: self.conversion_batch_complete(0, t, 0, 0))


    def process_single_file(self, work_item, use_ai, ai_client, ai_model, vision_model, custom_prompt):
        """Helper to process a single file in a worker thread."""
        item = work_item['id']
        input_path = work_item['input_path']
        filename = work_item['filename']
        output_path = work_item['output_path']
        
        # Check paused state (naive implementation for threads)
        while self.is_paused and not self.cancel_requested:
            import time
            time.sleep(0.5)

        if self.cancel_requested:
            return False, 0

        # Check budget
        if use_ai and not self.check_budget():
            self.safe_update(lambda: self.log(f"Budget limit reached for {filename}. Switching to local."))
            use_ai = False

        try:
             # Skip check was handled in main loop to update stats correctly, so we just proceed here
             
             self.safe_update(lambda: self.update_item_status(item, "Converting...", "-", "-"))
             
             def progress_cb(current, total, status=""):
                self.safe_update(lambda s=status, c=current, t=total: self.update_item_status(
                    item, 
                    f"{s}" if s else f"Page {c}/{t}",
                    "-", "-"
                ))

             # self.safe_update(lambda: self.log(f"Converting: {filename}")) # Too noisy for parallel

             try:
                 md_content, out_path, cost_info = self.converter.convert_file(
                    input_path,
                    output_path=output_path if self.output_format.get() == "markdown" else None,
                    extract_images=self.extract_images.get(),
                    detect_headers=True,
                    progress_callback=progress_cb,
                    use_ai=use_ai,
                    ai_client=ai_client,
                    ai_model=ai_model,
                    use_ocr=self.use_ocr.get(),
                    use_tables=self.use_tables.get(),
                    use_vision=self.use_vision.get(),
                    vision_model=vision_model if self.use_vision.get() else None,
                    custom_prompt=custom_prompt,
                    include_metadata=self.include_metadata.get(),
                    wrap_xml=self.wrap_xml.get(),
                    check_cancel=lambda: self.cancel_requested,
                    use_tika=self.config.get("use_tika", False),
                    use_tabula=self.config.get("use_tabula", False)
                )
             except Exception as e:
                 if use_ai and "API Error" in str(e):
                     self.safe_update(lambda: self.log(f"AI failed ({e}), falling back to local conversion..."))
                     # Fallback retry without AI
                     md_content, out_path, cost_info = self.converter.convert_file(
                        input_path,
                        output_path=output_path if self.output_format.get() == "markdown" else None,
                        extract_images=self.extract_images.get(),
                        detect_headers=True,
                        progress_callback=progress_cb,
                        use_ai=False, # Disable AI for retry
                        ai_client=None,
                        ai_model=None,
                        use_ocr=self.use_ocr.get(),
                        use_tables=self.use_tables.get(),
                        use_vision=False, # Disable Vision for retry
                        vision_model=None,
                        custom_prompt=custom_prompt,
                        include_metadata=self.include_metadata.get(),
                        wrap_xml=self.wrap_xml.get(),
                        check_cancel=lambda: self.cancel_requested,
                        use_tika=self.config.get("use_tika", False),
                        use_tabula=self.config.get("use_tabula", False)
                    )
                 else:
                     raise e

             # Handle formats
             fmt = self.output_format.get()
             if fmt == "html":
                 ExportManager.to_html(md_content, output_path)
             elif fmt == "docx":
                 ExportManager.to_docx(md_content, output_path)
             elif fmt == "txt":
                 ExportManager.to_txt(md_content, output_path)

             # Auto-vectorize for RAG if enabled
             if HAS_RAG and self.config.get('rag_settings', {}).get('auto_vectorize', True):
                 md_path = out_path if out_path else output_path
                 if md_path and os.path.exists(md_path):
                     self.vectorize_document(md_path, cost_info)

             cost = cost_info.get('cost', 0)
             
             # Track usage
             if use_ai and cost > 0:
                 self.safe_update(lambda: self.track_usage(cost, cost_info.get('input_tokens',0) + cost_info.get('output_tokens',0)))
             
             self.safe_update(lambda: self.update_item_status(item, "Done", f"${cost:.4f}", f"{cost_info.get('input_tokens',0)+cost_info.get('output_tokens',0)}"))
             self.safe_update(lambda: self.log(f"Converted: {filename}"))
             return True, cost

        except Exception as e:
            self.safe_update(lambda: self.log(f"Error converting {filename}: {str(e)}"))
            self.safe_update(lambda: self.update_item_status(item, "Failed", "Error", "-"))
            return False, 0

    def conversion_batch_complete(self, successful, failed, skipped, total_cost):
        """Handle batch conversion completion."""
        self.is_converting = False
        self.convert_btn.config(state="normal")
        self.pause_btn.config(state="disabled", text="Pause")
        self.cancel_btn.config(state="disabled")
        self.progress_label.config(text="Conversion complete")
        self.overall_progress_var.set(100)  # Ensure progress bar shows 100%

        if self.last_output_folder and Path(self.last_output_folder).exists():
            self.open_folder_btn.config(state="normal")

        total = successful + failed + skipped
        self.log(f"\n{'='*60}")
        self.log(f"CONVERSION COMPLETE!")
        self.log(f"{'='*60}")
        self.log(f"Successful: {successful}/{total}")
        if skipped > 0:
            self.log(f"Skipped: {skipped} file(s)")
        if total_cost > 0:
            self.log(f"Total API Cost: ${total_cost:.4f}")
        if failed > 0:
            self.log(f"Failed: {failed} file(s)")
        self.log(f"{'='*60}")

        self.status_var.set(f"Done: {successful} converted, {failed} failed, {skipped} skipped")
        self.update_usage_display()
        self.root.update_idletasks()  # Force final GUI update

        # Show completion dialog
        if failed == 0 and not self.cancel_requested:
            messagebox.showinfo("Success",
                f"Conversion Complete!\n\n"
                f"Successfully converted {successful} file(s).\n"
                f"Output folder: {self.last_output_folder}")
        elif self.cancel_requested:
            messagebox.showinfo("Cancelled",
                f"Conversion cancelled.\n\n"
                f"{successful} file(s) completed before cancellation.")
        else:
            messagebox.showwarning("Conversion Complete",
                f"Converted {successful} of {total} files.\n{failed} file(s) failed.\n\nCheck the log for details.")

    def on_closing(self):
        """Handle application shutdown."""
        try:
            self.cancel_requested = True
            
            # Save config/usage before exit
            if hasattr(self, 'config'):
                save_config(self.config)
            if hasattr(self, 'usage'):
                save_usage(self.usage)
                
            self.root.destroy()
            
            # Force kill any lingering threads/processes
            import os
            os._exit(0)
        except:
            sys.exit(0)

    def run(self):
        """Start the application."""
        self.root.mainloop()


def main():
    """Main entry point."""
    missing = check_dependencies()

    if missing:
        print("Missing required packages:")
        for pkg in missing:
            print(f"  - {pkg}")
        print("\nInstall them with:")
        print(f"  pip install {' '.join(missing)}")

        try:
            root = tk.Tk()
            root.withdraw()
            messagebox.showerror(
                "Missing Dependencies",
                f"Please install required packages:\n\npip install {' '.join(missing)}"
            )
            root.destroy()
        except:
            pass

        sys.exit(1)

    app = PDFToMarkdownApp()
    app.run()


if __name__ == "__main__":
    main()
